# BIO203A Lab 7: Endospore Staining — Traditional Lab Report (Popa STRICT v2)
# Author: George Vela
# References cite ONLY OpenStax Microbiology (Parker et al. 2016). CSE Name-Year style.
# SINGLE-SLIDE report: one Bacillus slant culture stained; one photograph documented.

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\User\Dropbox\Nu micro\lab reports\BIO203A_Lab7_Traditional_Report.docx"

STAIN_IMG_PATHS = [
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 7.JPG",
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 7.jpg",
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 7 endospore.JPG",
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 7 endospore stain.JPG",
]
STAIN_IMG = next((p for p in STAIN_IMG_PATHS if os.path.exists(p)), None)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
style.paragraph_format.space_after = Pt(0)

for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def add_page_number(footer_para):
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


for section in doc.sections:
    add_page_number(section.footer.paragraphs[0])


def add_paragraph(text_runs, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_indent=Inches(0.5), space_after=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.first_line_indent = first_indent
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = space_after
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'; run.font.size = Pt(fmt.get('size', 12))
        if fmt.get('bold'): run.bold = True
        if fmt.get('italic'): run.italic = True


def section_heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text); run.bold = True
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)


def subhead(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text); run.bold = True; run.italic = True
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)


def add_caption(text_runs):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(12)
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'; run.font.size = Pt(11)
        if fmt.get('bold'): run.bold = True
        if fmt.get('italic', True): run.italic = True


def add_image(path, width_inches=4.5):
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))


# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
run = p_title.add_run("Lab 7: Endospore Staining")
run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(16)

p_subtitle = doc.add_paragraph()
p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_subtitle.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
run = p_subtitle.add_run("Schaeffer-Fulton Endospore Staining of a Clostridium sporogenes Slant Culture")
run.italic = True; run.font.name = 'Times New Roman'; run.font.size = Pt(13)

for _ in range(3):
    doc.add_paragraph()

for line in [
    "George Vela",
    "BIO203A — Microbiology Laboratory",
    "Spring 2026",
    "Instructor: Dr. Radu Popa",
    "National University, Los Angeles Campus",
    "",
    "Staining Date: May 9, 2026",
    "Report Submitted: May 2026",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(line)
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SCOPE
# ════════════════════════════════════════════════════════════════════════════
section_heading("Scope")
add_paragraph([
    ("This laboratory exercise applied the Schaeffer-Fulton endospore-staining "
     "technique to a slant culture of ", {}),
    ("Clostridium sporogenes", {'italic': True}),
    (", an endospore-forming gram-positive obligate anaerobe provided by the "
     "instructor. The objective was to use a two-stain procedure to "
     "differentiate heat-resistant endospores from the surrounding vegetative "
     "cells in a single stained preparation and to observe the resulting "
     "differential stain at oil-immersion magnification.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
section_heading("Introduction")

add_paragraph([
    ("Endospores are structures produced within certain bacterial cells that "
     "essentially protect the bacterial genome in a dormant state when "
     "environmental conditions are unfavorable. Endospores allow some bacterial "
     "cells to survive long periods without food or water, as well as exposure "
     "to chemicals, extreme temperatures, and even radiation (Parker et al. 2016, "
     "§3.3). They are dormant, dehydrated, and metabolically inactive, in "
     "contrast to vegetative cells, which are capable of active growth and "
     "metabolism (Parker et al. 2016, §3.3).", {}),
])

add_paragraph([
    ("The process by which a vegetative cell transforms into an endospore is "
     "called sporulation, and it generally begins when nutrients become depleted "
     "or environmental conditions become otherwise unfavorable. Sporulation "
     "begins with the formation of a septum that divides the cell asymmetrically, "
     "separating a forespore that will form the core of the endospore. A cortex "
     "of calcium and dipicolinic acid is laid down around the forespore, a "
     "protein coat then forms around the cortex, and the endospore is released "
     "upon disintegration of the mother cell (Parker et al. 2016, §3.3).", {}),
])

add_paragraph([
    ("Endospores cannot be visualized by the Gram stain because the thick spore "
     "coat does not absorb crystal violet; in a Gram-stained preparation, "
     "endospores appear clear within the otherwise stained cell, and a special "
     "endospore stain is required to visualize them (Parker et al. 2016, §2.4 "
     "and §3.3). The Schaeffer-Fulton method is the most commonly used "
     "endospore-staining technique and uses heat to push the primary stain "
     "(malachite green) into the endospore. Washing with water decolorizes the "
     "cell, but the endospore retains the green stain. The cell is then "
     "counterstained pink with safranin (Parker et al. 2016, §2.4). The "
     "resulting image reveals the shape and location of endospores: green "
     "endospores appear either within pink vegetative cells or as separate "
     "green particles after the mother cell has disintegrated; if no endospores "
     "are present, only pink vegetative cells are visible (Parker et al. 2016, "
     "§2.4).", {}),
])

add_paragraph([
    ("Endospore-staining techniques are important for identifying ", {}),
    ("Bacillus", {'italic': True}),
    (", ", {}),
    ("Clostridium", {'italic': True}),
    (", and ", {}),
    ("Clostridioides", {'italic': True}),
    (", three genera of endospore-producing bacteria that contain clinically "
     "significant species (Parker et al. 2016, §2.4). Members of the genus ", {}),
    ("Bacillus", {'italic': True}),
    (" are large gram-positive bacilli that include aerobes or facultative "
     "anaerobes and form endospores (Parker et al. 2016, §4.4).", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# MATERIALS AND METHODS
# ════════════════════════════════════════════════════════════════════════════
section_heading("Materials and Methods")

subhead("Materials")
for item in [
    "Bacillus slant culture (provided by the instructor)",
    "Clean glass microscope slide",
    "Distilled water in a dropper",
    "Sterile inoculating loop",
    "Bunsen burner with striker (for slide heat-fixation)",
    "Clothespin (to hold slide during heat fixation)",
    "Malachite green stain (primary stain)",
    "Safranin stain (counterstain)",
    "Hot plate with beaker of water (for steaming the slide during staining)",
    "Small piece of absorbent paper (to retain malachite green on the smear)",
    "Forceps (to remove paper if it adheres to the slide)",
    "Staining tray, wash bottle of water, and waste beaker",
    "Bibulous paper for blotting",
    "Compound brightfield microscope with 100× oil immersion objective",
    "Type-A immersion oil",
    "Lens paper",
]:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(f"• {item}")
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)

subhead("Methods")

add_paragraph([
    ("On May 9, 2026, a smear was prepared on a labeled glass slide from a ", {}),
    ("Bacillus", {'italic': True}),
    (" slant culture provided by the instructor. A loopful of distilled water "
     "was placed on the slide, a small amount of culture was transferred from "
     "the slant into the water using a sterile inoculating loop, and the "
     "resulting suspension was spread thinly across a marked area of the "
     "slide. The slide was air-dried and then heat-fixed by passing it through "
     "the outer cone of the Bunsen flame two to three times with the smear "
     "side up.", {}),
])

add_paragraph([
    ("A beaker filled approximately one-third with water was placed on a hot "
     "plate and brought to a gentle simmer to produce steam without active "
     "boiling. The heat-fixed slide was placed across the rim of the beaker "
     "with the smear over the steam, and a small piece of absorbent paper was "
     "laid over the smear. Malachite green was applied to the paper to "
     "saturate it over the smear, and steaming was continued for five to seven "
     "minutes; additional malachite green was added as needed to keep the "
     "paper moist and prevent the slide from drying during steaming.", {}),
])

add_paragraph([
    ("After steaming, the slide was removed from the heat and allowed to cool "
     "briefly. The absorbent paper was rinsed off gently with distilled water, "
     "and the slide was washed with distilled water to decolorize any "
     "malachite green that had not been trapped in endospore coats. Safranin "
     "was applied as a counterstain to the smear and was left in contact for "
     "one minute, then rinsed with distilled water and blotted dry with "
     "bibulous paper.", {}),
])

add_paragraph([
    ("The stained slide was placed on the microscope stage, located at 4×, "
     "and brought into focus through the parfocal objective progression. A "
     "drop of immersion oil was applied directly to the smear and the 100× "
     "oil immersion objective was rotated into the oil for final observation "
     "at 1000× total magnification. A representative field was photographed "
     "through the ocular.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# RESULTS AND DISCUSSION
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
section_heading("Results and Discussion")

subhead("Photomicrograph")
add_image(STAIN_IMG, width_inches=4.5)
add_caption([
    ("Figure 1. ", {'bold': True, 'italic': True}),
    ("Schaeffer-Fulton endospore stain of a ", {'italic': True}),
    ("Bacillus", {'italic': True}),
    (" slant culture, photographed through the ocular at 1000× total "
     "magnification (100× oil immersion × 10× ocular). Numerous pink "
     "rod-shaped vegetative cells (safranin counterstain) are visible across "
     "the field, with smaller darker-stained structures distributed among "
     "them consistent with malachite-green-retaining endospores. The black "
     "needle-like feature in the lower portion of the field is the microscope "
     "stage pointer, not part of the specimen.", {'italic': True}),
])

subhead("Description of Observed Cells")
add_paragraph([
    ("The smear contained numerous rod-shaped cells that had taken up the "
     "safranin counterstain and appeared pink against the lighter background. "
     "Individual cells were visible as relatively uniform, straight bacilli. "
     "Among the pink vegetative cells, smaller, more darkly-staining structures "
     "were distributed within and around the rods; these were consistent in "
     "size and distribution with endospores that had retained the malachite "
     "green primary stain during the water decolorization step.", {}),
])

subhead("Interpretation")
add_paragraph([
    ("The contrast between the pink-stained vegetative cells and the darker "
     "endospore-consistent structures in the same preparation indicates that "
     "the Schaeffer-Fulton procedure achieved its intended differential effect "
     "on the smear. The Schaeffer-Fulton method uses heat to push the primary "
     "stain (malachite green) into the endospore; washing with water "
     "decolorizes the cell but the endospore retains the green stain, and the "
     "cell is then counterstained pink with safranin (Parker et al. 2016, "
     "§2.4). The resulting image — pink vegetative cells with green endospores "
     "within or beside them — is the expected outcome of a successful "
     "Schaeffer-Fulton stain, and is what was observed on this slide.", {}),
])

add_paragraph([
    ("Vegetative cells are sensitive to extreme temperatures and radiation and "
     "have normal water content and enzymatic activity, whereas endospores are "
     "resistant to extreme temperatures and radiation, are dehydrated, and "
     "have no metabolic activity (Parker et al. 2016, §3.3). The thick spore "
     "coat that confers this resistance is also responsible for the staining "
     "behavior observed here: endospores do not absorb the Gram stain and "
     "require special endospore-staining techniques, and even with the "
     "Schaeffer-Fulton method, heat must be applied during the malachite green "
     "step to drive the primary stain through the spore coat (Parker et al. "
     "2016, §2.4 and §3.3).", {}),
])

add_paragraph([
    ("Three of the four major endospore-producing bacterial genera — ", {}),
    ("Bacillus", {'italic': True}),
    (", ", {}),
    ("Clostridium", {'italic': True}),
    (", and ", {}),
    ("Clostridioides", {'italic': True}),
    (" — contain clinically significant species, and endospore staining is a "
     "diagnostic technique used to identify suspected isolates of these "
     "groups (Parker et al. 2016, §2.4). Within ", {}),
    ("Bacillus", {'italic': True}),
    (", important pathogens include ", {}),
    ("B. anthracis", {'italic': True}),
    (" (anthrax) and ", {}),
    ("B. cereus", {'italic': True}),
    (" (a food-poisoning agent) (Parker et al. 2016, §4.4). The "
     "endospore-staining technique practiced in this exercise is therefore a "
     "diagnostic foundation rather than an end in itself.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ════════════════════════════════════════════════════════════════════════════
section_heading("Conclusion")

add_paragraph([
    ("The Schaeffer-Fulton endospore-staining procedure was applied to a slant "
     "culture of a ", {}),
    ("Bacillus", {'italic': True}),
    (" species and produced a stained preparation in which both pink-stained "
     "vegetative cells and darker-stained endospore-consistent structures were "
     "visible in the same microscope field at 1000× total magnification. The "
     "outcome is consistent with the staining mechanism described in OpenStax "
     "§2.4 — endospores retain malachite green after the water rinse while "
     "vegetative cells are decolorized and take up safranin — and with the "
     "endospore-formation biology described in §3.3.", {}),
])

add_paragraph([
    ("Endospore staining is a key diagnostic technique for identifying the "
     "clinically relevant genera ", {}),
    ("Bacillus", {'italic': True}),
    (", ", {}),
    ("Clostridium", {'italic': True}),
    (", and ", {}),
    ("Clostridioides", {'italic': True}),
    (" (Parker et al. 2016, §2.4). The procedure practiced in this exercise "
     "will be applied to gram-positive rod isolates from the library plate in "
     "the Antibiotic Discovery Project to determine whether any of the soil "
     "isolates are endospore-formers.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════════════
section_heading("References")


def add_reference(text_runs):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(0)
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'; run.font.size = Pt(12)


add_reference([
    ("Parker N, Schneegurt M, Tu A-HT, Lister P, Forster BM. 2016. Microbiology. "
     "Houston (TX): OpenStax. Available from: "
     "https://openstax.org/details/books/microbiology", {}),
])

doc.save(OUTPUT)
print(f"Done -> {OUTPUT}")
print(f"Endospore stain photo: {'EMBEDDED' if STAIN_IMG else 'MISSING'}")
if STAIN_IMG: print(f"  {STAIN_IMG}")
