# BIO203A Lab 1: Microscopy — Traditional Lab Report (Word .docx)
# Formatted to standard college biology lab report conventions:
#   - 12pt Times New Roman, double-spaced, 1" margins
#   - Title page, then body with bold left-aligned section headers
#   - Past-tense passive in Materials/Methods, past tense in Results,
#     present tense for established facts in Introduction
#   - Scientific names italicized (Trypanosoma)
#   - Figures numbered with captions BELOW the image (Figure 1, 2, ...)
#   - CSE Name-Year reference style (standard for biology)
# References cite ONLY OpenStax + Tiny Earth manual + Popa lecture slides.

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\User\Dropbox\Nu micro\lab reports\BIO203A_Lab1_Traditional_Report.docx"
DRAWING_PATH = r"C:\Users\User\Dropbox\Nu micro\lab1_drawing.png"
FIELD_SKETCH_PATH = r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 1 image.png"

doc = Document()

# ── Document-wide format defaults ────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
style.paragraph_format.space_after = Pt(0)

# 1-inch margins everywhere
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


# Page numbers in footer (right-aligned)
def add_page_number(footer_para):
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


for section in doc.sections:
    add_page_number(section.footer.paragraphs[0])


# ── Helper functions ─────────────────────────────────────────────────────────
def add_paragraph(text_runs, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_indent=Inches(0.5), space_after=Pt(0)):
    """text_runs is a list of (text, formatting) where formatting is a dict
    with optional keys: bold, italic, size, underline."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.first_line_indent = first_indent
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = space_after
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(fmt.get('size', 12))
        if fmt.get('bold'):
            run.bold = True
        if fmt.get('italic'):
            run.italic = True
        if fmt.get('underline'):
            run.underline = True
    return p


def section_heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_caption(text_runs):
    """Figure caption: italic, centered, single-spaced (per scientific writing convention)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(12)
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        if fmt.get('bold'):
            run.bold = True
        if fmt.get('italic', True):  # default italic for captions
            run.italic = True
    return p


def add_image(path, width_inches=5.0):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))


# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
# Push title to roughly center vertically with blank lines
for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
run = p_title.add_run("Lab 1: Microscopy")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(16)

p_subtitle = doc.add_paragraph()
p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_subtitle.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
run = p_subtitle.add_run("Operation of the Compound Brightfield Microscope and Observation of a Stained Human-Tissue Slide")
run.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(13)

# Spacing
for _ in range(3):
    doc.add_paragraph()

for line in [
    "George Vela",
    "BIO203A — Microbiology Laboratory",
    "Spring 2026",
    "Instructor: Dr. Radu Popa",
    "National University, Los Angeles Campus",
    "",
    "Lab Date: April 28, 2026",
    "Report Submitted: May 2026",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(line)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. SCOPE
# ════════════════════════════════════════════════════════════════════════════
section_heading("Scope")
add_paragraph([
    ("This laboratory exercise introduced the use of the compound brightfield light "
     "microscope and applied it to the observation of a prepared, stained slide of "
     "human tissue. The investigation focused on the operation of the microscope at "
     "increasing magnifications, the relationships between magnification, field of "
     "view, and resolution, and the description of the morphology of the cells "
     "observed on the slide.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# 2. INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
section_heading("Introduction")

add_paragraph([
    ("An object must measure about 100 µm to be visible without a microscope, but "
     "most microorganisms are many times smaller than that. A typical animal cell "
     "measures roughly 10 µm across, bacterial cells are typically about 1 µm, and "
     "viruses can be 10 times smaller than bacteria (Parker et al. 2016, §1.3).", {}),
])

add_paragraph([
    ("The brightfield microscope is a compound microscope with two or more lenses "
     "that produce a dark image on a bright background (Parker et al. 2016, §2.3). "
     "Each eyepiece contains a lens called an ocular lens, and the ocular lenses "
     "typically magnify images 10×. At the other end of the body tube are a set of "
     "objective lenses on a rotating nosepiece, with magnification of these "
     "objective lenses typically ranging from 4× to 100× (Parker et al. 2016, §2.3). "
     "Total magnification is the product of the ocular magnification times the "
     "objective magnification (Parker et al. 2016, §2.3).", {}),
])

add_paragraph([
    ("In a brightfield microscope, light is provided by an illuminator below the "
     "stage, passes up through the condenser lens which focuses light rays on the "
     "specimen, and the amount of light striking the specimen can be adjusted by "
     "opening or closing a diaphragm between the condenser and the specimen "
     "(Parker et al. 2016, §2.3). The coarse focusing knob is used for large-scale "
     "movements with the 4× and 10× objective lenses; the fine focusing knob is "
     "used for small-scale movements, especially with 40× or 100× objective lenses "
     "(Parker et al. 2016, §2.3).", {}),
])

add_paragraph([
    ("At very high magnifications, resolution may be compromised when light passes "
     "through the small amount of air between the specimen and the lens, due to "
     "the large difference between the refractive indices of air and glass which "
     "scatters the light rays before they can be focused by the lens (Parker et al. 2016, "
     "§2.3). To solve this problem, a drop of immersion oil is used at the 100× "
     "objective (Parker et al. 2016, §2.3).", {}),
])

add_paragraph([
    ("Microorganisms include several major groups; protozoa are unicellular protists "
     "that may move using cilia or flagella, and some are parasitic, surviving by "
     "extracting nutrients from a host organism (Parker et al. 2016, §1.3). "
     "Identification of microscopic organisms based on observed morphology alone is "
     "tentative; confirmed identification typically requires staining, biochemical "
     "testing, or molecular methods that go beyond the scope of this introductory "
     "microscopy exercise.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# 3. MATERIALS AND METHODS
# ════════════════════════════════════════════════════════════════════════════
section_heading("Materials and Methods")

# Materials sub-heading
p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
p.paragraph_format.space_before = Pt(6)
run = p.add_run("Materials")
run.bold = True
run.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

materials_list = [
    "Compound binocular brightfield microscope with 4×, 10×, 40×, and 100× oil immersion objectives, 10× oculars, and a field number of 22 mm",
    "Pre-prepared, fixed and stained microscope slides of human tissue (containing red blood cells and additional smaller cells of unknown identity)",
    "Type-A immersion oil",
    "Lens paper",
    "Permanent marker for slide identification",
]
for item in materials_list:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(f"• {item}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# Methods sub-heading
p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
p.paragraph_format.space_before = Pt(6)
run = p.add_run("Methods")
run.bold = True
run.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

add_paragraph([
    ("The microscope was carried to the bench using both hands — one supporting the "
     "base, one grasping the arm — and the power cord was placed beside the unit "
     "rather than wrapped around the stage. The 4× scanning objective was rotated "
     "into position and the stage was raised to its lowest setting using the "
     "coarse-adjustment knob. The illuminator was switched on and the iris diaphragm "
     "of the condenser was partially closed to increase contrast at low magnification.", {}),
])

add_paragraph([
    ("The prepared slide was secured on the stage. Initial focus was achieved at 4× "
     "and the specimen was centered in the field of view. The nosepiece was then "
     "rotated to the 10× objective, which required only minor refocusing using the "
     "fine-adjustment knob, confirming the parfocal property of the lens system. "
     "The 40× high-dry objective was engaged next and the iris diaphragm was opened "
     "wider to compensate for the smaller field of view and reduced light. A "
     "general sketch of cell size and shape was recorded at this magnification.", {}),
])

add_paragraph([
    ("For the highest-resolution observation, the 40× objective was rotated out of "
     "the optical path. A single drop of immersion oil was placed directly on the "
     "area of the slide being observed, and the 100× oil immersion objective was "
     "rotated into the oil. Fine focus was adjusted carefully to bring the "
     "specimen into sharp resolution. Cell shape, internal organelles, and "
     "identifying features were recorded by hand sketch.", {}),
])

add_paragraph([
    ("Upon completion, the slide was removed without rotating the stage downward "
     "against the oil objective. The 100× lens was wiped clean with lens paper, "
     "the stage was cleaned of residual oil, the nosepiece was returned to the 4× "
     "objective, the stage was lowered to its furthest position from the lenses, "
     "and the microscope was returned to the storage cabinet with the arm pointing "
     "outward.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# 4. RESULTS AND DISCUSSION
# ════════════════════════════════════════════════════════════════════════════
section_heading("Results and Discussion")

# Organism observed
p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
p.paragraph_format.space_before = Pt(6)
run = p.add_run("Organism Observed")
run.bold = True
run.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

add_paragraph([
    ("The slide contained two distinct cell types observed in the same field. The "
     "dominant cells were biconcave, anucleate discs approximately 7–8 µm in "
     "diameter — morphology consistent with red blood cells. Among them, smaller "
     "and less numerous slender elongated cells were observed; each had a clear "
     "central nucleus, a smaller darker-stained granule located adjacent to the "
     "nucleus, and a single long flagellum trailing from one end of the cell. "
     "These features are consistent with a flagellated single-celled organism, but "
     "definitive identification was not made during this exercise.", {}),
])

# Figure 1
p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
p.paragraph_format.space_before = Pt(6)
run = p.add_run("Sketch of Specimens at 400× Total Magnification (40× objective × 10× ocular)")
run.bold = True
run.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

add_image(DRAWING_PATH, width_inches=5.5)
add_caption([
    ("Figure 1. ", {'bold': True, 'italic': True}),
    ("Hand sketch of the slide at 400× total magnification. Left: an anucleate "
     "biconcave disc consistent with a red blood cell. Right: a smaller elongated "
     "cell with a single nucleus (labeled N), a smaller darker granule near the "
     "nucleus (labeled K), and a single flagellum extending from one end. Labels "
     "applied during the lab session reflect a tentative working identification "
     "based on observed morphology only.", {'italic': True}),
])

# Figure 2
add_image(FIELD_SKETCH_PATH, width_inches=5.0)
add_caption([
    ("Figure 2. ", {'bold': True, 'italic': True}),
    ("Field-notebook sketch made during the lab session. Center: a larger host cell "
     "appearing to contain multiple smaller intracellular structures. Right: two "
     "smaller elongated cells with a flagellum extending from one end, located "
     "outside the host cell. Definitive identification of these organisms was "
     "not made during the exercise.", {'italic': True}),
])

# FOV calculations
p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
p.paragraph_format.space_before = Pt(6)
run = p.add_run("Field of View and Cell Count Calculations")
run.bold = True
run.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

add_paragraph([
    ("At the 40× objective with a field number (FN) of 22 mm and a 10× ocular, "
     "the field of view was calculated as FN ÷ total magnification = 22 mm ÷ 400 = "
     "0.055 mm = ", {}),
    ("55 µm", {'bold': True}),
    (". Using this FOV, approximately 27 ", {}),
    ("Bacillus", {'italic': True}),
    (" cells (2 µm long) or approximately 5–6 yeast cells (10 µm long) would fit "
     "across the diameter of the field. This calculation illustrates the inverse "
     "relationship between magnification and field of view: as magnification "
     "increases, the visible field shrinks and fewer cells fit across it. The same "
     "relationship explains why locating a specimen begins at low power (4× scanning, "
     "where the field is widest) before progressing to higher objectives.", {}),
])

# Cell count summary table
table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

hdr = table.rows[0].cells
for i, txt in enumerate(['Cell type', 'Length (µm)', 'Cells across 55 µm FOV']):
    hdr[i].text = ''
    p = hdr[i].paragraphs[0]
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(txt)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

rows = [
    ['Bacillus (rod-shaped bacterium)', '2', '55 ÷ 2 ≈ 27 cells'],
    ['Yeast (Saccharomyces)', '10', '55 ÷ 10 ≈ 5.5 cells'],
]
for i, row_data in enumerate(rows, start=1):
    cells = table.rows[i].cells
    for j, txt in enumerate(row_data):
        cells[j].text = ''
        p = cells[j].paragraphs[0]
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(txt)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        if j == 0 and 'Bacillus' in txt:
            # italicize Bacillus and Saccharomyces
            pass

# Italicize Bacillus and Saccharomyces in cells - manual fix
for i, row_data in enumerate(rows, start=1):
    cell = table.rows[i].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if 'Bacillus' in row_data[0]:
        run1 = p.add_run("Bacillus")
        run1.italic = True
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(11)
        run2 = p.add_run(" (rod-shaped bacterium)")
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(11)
    elif 'Yeast' in row_data[0]:
        run1 = p.add_run("Yeast (")
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(11)
        run2 = p.add_run("Saccharomyces")
        run2.italic = True
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(11)
        run3 = p.add_run(")")
        run3.font.name = 'Times New Roman'
        run3.font.size = Pt(11)

# Spacer after table
p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

# Discussion paragraph (replaces the old handout-style Q&A block)
p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
p.paragraph_format.space_before = Pt(6)
run = p.add_run("Discussion")
run.bold = True
run.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

add_paragraph([
    ("The progression from low- to high-power objectives confirmed the inverse "
     "relationship between magnification and field of view: as total magnification "
     "increased, the visible field shrank and proportionally fewer cells were "
     "visible at one time. The color of the cells did not change between objectives, "
     "consistent with stain-bound color being a property of the specimen rather "
     "than of the optical system. At high magnification the 100× oil immersion "
     "objective is closest to the slide, and so only the fine focusing knob — "
     "intended for small-scale movements with 40× and 100× objectives — was used "
     "during high-power observation; the coarse focusing knob is reserved for "
     "large-scale movements at 4× and 10× (Parker et al. 2016, §2.3). Light delivery "
     "at higher magnification was managed through the diaphragm between the "
     "condenser and the specimen, supplemented by the illuminator's intensity "
     "control, and at 100× the immersion oil itself reduced refractive losses at "
     "the glass-air interface (Parker et al. 2016, §2.3).", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# 5. CONCLUSION
# ════════════════════════════════════════════════════════════════════════════
section_heading("Conclusion")

add_paragraph([
    ("The objectives of Lab 1 were met. The compound brightfield microscope was "
     "operated through the progression from the 4× to the 100× oil immersion "
     "objective. Total magnification was calculated as the product of the ocular "
     "magnification times the objective magnification (Parker et al. 2016, §2.3) "
     "and applied to a practical cell-counting calculation. Two distinct cell types "
     "were observed on the prepared human-tissue slide: biconcave anucleate discs "
     "consistent with red blood cells, and smaller elongated flagellated cells whose "
     "morphology was described but whose species-level identity was not confirmed.", {}),
])

add_paragraph([
    ("The skills practiced in this laboratory — slide handling, objective lens "
     "progression, oil immersion technique, and description of cell morphology — "
     "will be applied in subsequent labs in which microorganisms isolated from "
     "environmental samples are examined microscopically.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# 6. REFERENCES (CSE Name-Year style, alphabetical)
# ════════════════════════════════════════════════════════════════════════════
section_heading("References")

# Hanging indent for CSE-style references
def add_reference(text_runs):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(0)
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if fmt.get('italic'):
            run.italic = True
        if fmt.get('bold'):
            run.bold = True


add_reference([
    ("Parker N, Schneegurt M, Tu A-HT, Lister P, Forster BM. 2016. Microbiology. "
     "Houston (TX): OpenStax. Available from: "
     "https://openstax.org/details/books/microbiology", {}),
])

# Save
doc.save(OUTPUT)
print(f"Done -> {OUTPUT}")
