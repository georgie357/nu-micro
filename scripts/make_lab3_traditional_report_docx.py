# BIO203A Lab 3: Pick and Patch of Solid Colonies — Traditional Lab Report (Popa format, STRICT v2)
# Author: George Vela
# References cite ONLY OpenStax Microbiology (Parker et al. 2016). CSE Name-Year style.

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\User\Dropbox\Nu micro\lab reports\BIO203A_Lab3_Traditional_Report.docx"

LIBRARY_PATHS = [
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 3 library plate.jpg",
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 3 library plate.png",
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 3 library plate.JPG",
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 3 library plate.jpg.JPG",
]
LABEL_PATHS = [
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 3 library plate label.jpg",
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 3 library plate label.png",
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 3 library plate label.JPG",
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 3 library plate label.jpg.JPG",
]
LIBRARY_PATH = next((p for p in LIBRARY_PATHS if os.path.exists(p)), None)
LABEL_PATH = next((p for p in LABEL_PATHS if os.path.exists(p)), None)

# Source plate from Lab 2 (already saved)
SOURCE_PLATE = r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 2 plate.jpg.JPG"

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
style.paragraph_format.space_after = Pt(0)

for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def add_page_number(footer_para):
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


for section in doc.sections:
    add_page_number(section.footer.paragraphs[0])


def add_paragraph(text_runs, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_indent=Inches(0.5), space_after=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.first_line_indent = first_indent
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = space_after
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(fmt.get('size', 12))
        if fmt.get('bold'): run.bold = True
        if fmt.get('italic'): run.italic = True
    return p


def section_heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


def subhead(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text)
    run.bold = True; run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


def add_caption(text_runs):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(12)
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        if fmt.get('bold'): run.bold = True
        if fmt.get('italic', True): run.italic = True


def add_image(path, width_inches=4.5):
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run("[Photo placeholder — save image to expected path and re-run]")
        run.font.name = 'Times New Roman'; run.font.size = Pt(11); run.italic = True


# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
run = p_title.add_run("Lab 3: Pick and Patch of Solid Colonies")
run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(16)

p_subtitle = doc.add_paragraph()
p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_subtitle.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
run = p_subtitle.add_run("Construction of a Bacterial Library Plate from Backyard Garden Soil for Antibiotic Screening")
run.italic = True; run.font.name = 'Times New Roman'; run.font.size = Pt(13)

for _ in range(3):
    doc.add_paragraph()

for line in [
    "George Vela",
    "BIO203A — Microbiology Laboratory",
    "Spring 2026",
    "Instructor: Dr. Radu Popa",
    "National University, Los Angeles Campus",
    "",
    "Source Plate Date: April 30, 2026",
    "Pick and Patch Date: May 5, 2026",
    "Report Submitted: May 2026",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(line)
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SCOPE
# ════════════════════════════════════════════════════════════════════════════
section_heading("Scope")
add_paragraph([
    ("This laboratory exercise applied the pick-and-patch technique to isolate "
     "morphologically distinct bacterial colonies from a soil dilution plate and to "
     "assemble them into a single library plate of separated, traceable patches. The "
     "library plate serves as the source of pure isolates for downstream antibiotic "
     "screening in the semester-long Antibiotic Discovery Project.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
section_heading("Introduction")

add_paragraph([
    ("Soil bacteria have historically been the most productive source of clinically "
     "useful antimicrobials. Selman Waksman, a soil microbiologist at Rutgers "
     "University, led a research team that discovered several antimicrobials including "
     "actinomycin, streptomycin, and neomycin; the discoveries stemmed from Waksman's "
     "study of fungi and the Actinobacteria, including soil bacteria in the genus ", {}),
    ("Streptomyces", {'italic': True}),
    (", known for their natural production of a wide variety of antimicrobials "
     "(Parker et al. 2016, §14.1). Soil remains an excellent reservoir for the "
     "discovery of novel antimicrobial agents, and some researchers argue that the "
     "soil microbiome has not yet been fully exploited as a source of new drugs "
     "(Parker et al. 2016, §14.7). The genus ", {}),
    ("Streptomyces", {'italic': True}),
    (" alone accounts for more than two-thirds of clinically useful antibiotics "
     "(Parker et al. 2016, §4.4).", {}),
])

add_paragraph([
    ("Discovery of antimicrobial compounds from environmental microbes requires the "
     "isolation of pure bacterial cultures from mixed environmental samples. A "
     "viable plate count from a soil dilution typically yields plates with multiple "
     "morphologically distinct colonies, each arising from a single cell or a small "
     "cluster of cells (Parker et al. 2016, §9.1). Selecting individual colonies and "
     "transferring them to a fresh, well-organized plate produces a catalog of "
     "isolates from the same original sample — a so-called library plate — in which "
     "each patch is a candidate pure culture for subsequent antimicrobial-activity "
     "testing.", {}),
])

add_paragraph([
    ("Colonies most likely to produce antimicrobials are those whose appearance is "
     "consistent with secondary-metabolite-producing soil bacteria. Pigmented "
     "colonies and colonies with filamentous, dry, or chalky morphology are typical "
     "indicators of ", {}),
    ("Streptomyces", {'italic': True}),
    (" and related Actinobacteria, which give soil its characteristic earthy odor "
     "and are aerobic, spore-forming, filamentous bacteria (Parker et al. 2016, §4.4). "
     "These morphologies were therefore prioritized during colony selection for the "
     "library plate constructed in this exercise.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# MATERIALS AND METHODS
# ════════════════════════════════════════════════════════════════════════════
section_heading("Materials and Methods")

subhead("Materials")
for item in [
    "Source dilution plate: TSA plate of 10⁻³ dilution of backyard garden soil (prepared 4/30/2026; 22 countable colonies in 6 morphological types)",
    "Fresh Tryptic Soy Agar (TSA) library plate, pre-poured",
    "Sterile wooden toothpicks (one per colony picked)",
    "Permanent marker for plate labeling and colony marking",
    "Paper grid (used as backing under the library plate to orient the patch grid)",
    "Bunsen burner (for sterile-technique work area)",
    "Incubator (room temperature, ~25 °C)",
]:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(f"• {item}")
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)

subhead("Methods")

add_paragraph([
    ("On May 5, 2026, candidate colonies were selected from the source TSA plate "
     "that had been prepared from a 10⁻³ soil dilution on April 30, 2026. Selection "
     "prioritized morphological uniqueness; pigmented colonies and colonies with "
     "distinct surface textures were preferred. A fresh TSA plate was labeled with "
     "investigator initials, the date, and the designation \"LP\" (library plate), "
     "and the back of the plate was divided with permanent marker into a grid of "
     "numbered squares. A small orientation line was drawn at the edge of the plate.", {}),
])

add_paragraph([
    ("To pick each colony, a single sterile wooden toothpick was used to lightly "
     "touch the surface of a marked colony on the source plate. The toothpick was "
     "then used to patch the picked material onto a single grid square of the "
     "library plate by gentle zigzag motion. A new sterile toothpick was used for "
     "each colony to prevent cross-contamination of patches. Each square of the "
     "library plate received one colony, and care was taken to keep patches within "
     "the boundaries of their squares so that adjacent patches did not touch.", {}),
])

add_paragraph([
    ("A total of seven colonies were picked and patched. Six were taken from the "
     "investigator's own source plate, and one additional colony was picked from a "
     "different student's source plate to broaden the diversity of isolates on the "
     "library plate. After all patches were applied, the library plate was inverted "
     "and incubated at room temperature (~25 °C) to allow colony regrowth at each "
     "patch site.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# RESULTS AND DISCUSSION
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
section_heading("Results and Discussion")

subhead("Source Plate")
add_paragraph([
    ("The source plate (Figure 1) was the 10⁻³ TSA plate prepared in Lab 2 from the "
     "backyard garden soil sample, containing 22 well-separated colonies grouped into "
     "six morphologically distinguishable types: large bright yellow colonies, "
     "medium and small yellow colonies, white/cream pinpoint colonies, a single "
     "white wrinkled colony, and two white filamentous-textured colonies. Colonies "
     "for the library plate were selected to represent this morphological diversity, "
     "with preference given to pigmented and filamentous types.", {}),
])

add_image(SOURCE_PLATE, width_inches=4.0)
add_caption([
    ("Figure 1. ", {'bold': True, 'italic': True}),
    ("Source plate (TSA, 10⁻³ dilution, plated 4/30/2026) from which colonies were "
     "selected for the library plate. Multiple yellow colonies of varying sizes "
     "are visible along with one white filamentous-textured colony and white "
     "pinpoint colonies.", {'italic': True}),
])

subhead("Library Plate after Incubation")
add_paragraph([
    ("The completed library plate (Figure 2) contained seven patched colonies, each "
     "in its own numbered grid square. After incubation, regrowth was visible at "
     "each patched location, with the patches preserving the color and approximate "
     "surface texture of the source colonies from which they were taken (Table 1). "
     "Pigmented yellow patches were the most numerous, consistent with the dominance "
     "of yellow colonies on the source plate. One patch showed a white "
     "filamentous-textured colony of the type morphologically suggestive of ", {}),
    ("Streptomyces", {'italic': True}),
    (" (Parker et al. 2016, §4.4), and is a particularly promising candidate for "
     "antimicrobial screening.", {}),
])

add_image(LIBRARY_PATH, width_inches=4.5)
add_caption([
    ("Figure 2. ", {'bold': True, 'italic': True}),
    ("Library plate after incubation, showing seven patched colonies in separate "
     "grid squares on TSA. Patches retained the color and approximate morphology "
     "of the source colonies from which they were taken.", {'italic': True}),
])

if LABEL_PATH:
    add_image(LABEL_PATH, width_inches=3.5)
    add_caption([
        ("Figure 3. ", {'bold': True, 'italic': True}),
        ("Close-up of the library plate label showing plating volume (75 µL), source "
         "dilution (10⁻³), date (5/5/2026), investigator initials (GV), and "
         "designation \"TSA chrome libry\" indicating the TSA library plate for "
         "the chromogenic (pigmented) isolates.", {'italic': True}),
    ])

subhead("Patch Inventory")
add_paragraph([
    ("Patches were numbered to correspond to grid position on the library plate. "
     "Morphological description of each patch is summarized in Table 1.", {}),
])

# Table
table = doc.add_table(rows=8, cols=4)
table.style = 'Table Grid'
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

headers = ['Patch #', 'Color', 'Surface / Edge', 'Source']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]; cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(h); run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(11)

patches = [
    ['1', 'White / cream', 'Cluster, slightly rough', 'Another student\'s source plate'],
    ['2', 'Yellow', 'Smooth, raised, circular', 'Own 10⁻³ plate (type 2)'],
    ['3', 'White / cream', 'Cluster, rough surface', 'Own 10⁻³ plate (type 4)'],
    ['4', 'Yellow', 'Smooth, small, round', 'Own 10⁻³ plate (type 3)'],
    ['5', 'Yellow', 'Large, irregular outline (kidney shape)', 'Own 10⁻³ plate (type 1)'],
    ['6', 'White / cream', 'Small, clustered', 'Own 10⁻³ plate (type 4)'],
    ['7', 'Yellow', 'Medium, smooth', 'Own 10⁻³ plate (type 2)'],
]
for i, row in enumerate(patches, start=1):
    for j, val in enumerate(row):
        cell = table.rows[i].cells[j]; cell.text = ''
        p = cell.paragraphs[0]; p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(val); run.font.name = 'Times New Roman'; run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(12)
run = p.add_run("Table 1. ")
run.bold = True; run.italic = True; run.font.name = 'Times New Roman'; run.font.size = Pt(11)
run = p.add_run(
    "Morphological description of the seven patches on the library plate after incubation. "
    "Source-plate \"type\" numbers refer to the six morphological groups identified in the Lab 2 plate inventory."
)
run.italic = True; run.font.name = 'Times New Roman'; run.font.size = Pt(11)

subhead("Significance for the Antibiotic Discovery Project")
add_paragraph([
    ("The library plate provides a small but morphologically diverse panel of "
     "candidate isolates from a single environmental source. Yellow pigmentation in "
     "soil bacteria is commonly associated with carotenoid and other secondary-metabolite "
     "production, and dry filamentous-textured colonies are characteristic of "
     "Actinobacteria including the genus ", {}),
    ("Streptomyces", {'italic': True}),
    (", which is responsible for more than two-thirds of clinically useful antibiotics "
     "(Parker et al. 2016, §4.4). Soil is described as an excellent reservoir for the "
     "discovery of novel antimicrobial agents (Parker et al. 2016, §14.7), and the "
     "panel of patches captured on this library plate will serve as the starting "
     "material for the antimicrobial-overlay and cross-streak assays scheduled for "
     "later weeks of the course.", {}),
])

add_paragraph([
    ("Two limitations of the library plate should be noted. First, the total number "
     "of patches (seven) is below the 10–12 unique colonies typically recommended "
     "for a library plate, reducing the probability of capturing rare antibiotic "
     "producers from the sample. Second, one patch was sourced from another "
     "student's plate and therefore originates from a different soil sample with "
     "an unknown collection environment; its identity and inclusion in any "
     "subsequent screening assay must be documented separately so that downstream "
     "results can be traced to the correct soil source.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ════════════════════════════════════════════════════════════════════════════
section_heading("Conclusion")

add_paragraph([
    ("A library plate of seven morphologically distinct bacterial patches was "
     "successfully constructed by the pick-and-patch technique from the 10⁻³ TSA "
     "source plate of the Lab 2 backyard garden soil sample, with one additional "
     "patch contributed from a different student's source plate. Patch morphologies "
     "after incubation preserved the color and texture of the source colonies and "
     "included several pigmented yellow colonies as well as one filamentous-textured "
     "colony of the kind morphologically suggestive of ", {}),
    ("Streptomyces", {'italic': True}),
    (" (Parker et al. 2016, §4.4).", {}),
])

add_paragraph([
    ("The library plate is now a working catalog of pure-culture candidates from the "
     "sampled soil environment and is ready to serve as the source of isolates for "
     "the antimicrobial-screening exercises in subsequent weeks of the Antibiotic "
     "Discovery Project.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════════════
section_heading("References")


def add_reference(text_runs):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(0)
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'; run.font.size = Pt(12)
        if fmt.get('italic'): run.italic = True


add_reference([
    ("Parker N, Schneegurt M, Tu A-HT, Lister P, Forster BM. 2016. Microbiology. "
     "Houston (TX): OpenStax. Available from: "
     "https://openstax.org/details/books/microbiology", {}),
])

doc.save(OUTPUT)
print(f"Done -> {OUTPUT}")
if LIBRARY_PATH:
    print(f"Library plate photo embedded: {LIBRARY_PATH}")
else:
    print(f"NO library plate photo found. Expected paths:")
    for p in LIBRARY_PATHS:
        print(f"  {p}")
if LABEL_PATH:
    print(f"Label close-up embedded: {LABEL_PATH}")
else:
    print(f"NO label close-up found (optional). Expected paths:")
    for p in LABEL_PATHS:
        print(f"  {p}")
