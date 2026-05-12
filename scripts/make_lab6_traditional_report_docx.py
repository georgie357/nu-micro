# BIO203A Lab 6: Gram Staining — Traditional Lab Report (Popa STRICT v2)
# Author: George Vela
# References cite ONLY OpenStax Microbiology (Parker et al. 2016). CSE Name-Year style.
# NOTE: Image paths are placeholders. Save photos to:
#   C:\Users\User\Dropbox\Nu micro\lab reports\lab 6 gram bacillus.jpg
#   C:\Users\User\Dropbox\Nu micro\lab reports\lab 6 gram ecoli.jpg
# Then re-run this script.

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\User\Dropbox\Nu micro\lab reports\BIO203A_Lab6_Traditional_Report.docx"

MIXED_IMG_PATHS = [
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 6 variable stain.JPG",
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 6 variable stain.jpg",
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 6 mixed stain.JPG",
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 6 mixed stain.jpg",
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 6 gram mixed.JPG",
]
MIXED_IMG = next((p for p in MIXED_IMG_PATHS if os.path.exists(p)), None)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
style.paragraph_format.space_after = Pt(0)

for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def add_page_number(footer_para):
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


for section in doc.sections:
    add_page_number(section.footer.paragraphs[0])


def add_paragraph(text_runs, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_indent=Inches(0.5), space_after=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.first_line_indent = first_indent
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = space_after
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'; run.font.size = Pt(fmt.get('size', 12))
        if fmt.get('bold'): run.bold = True
        if fmt.get('italic'): run.italic = True


def section_heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text); run.bold = True
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)


def subhead(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text); run.bold = True; run.italic = True
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)


def add_caption(text_runs):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(12)
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'; run.font.size = Pt(11)
        if fmt.get('bold'): run.bold = True
        if fmt.get('italic', True): run.italic = True


def add_image(path, width_inches=4.5):
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run("[Photo placeholder — save image to expected path and re-run script]")
        run.font.name = 'Times New Roman'; run.font.size = Pt(11); run.italic = True


# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
run = p_title.add_run("Lab 6: Gram Staining")
run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(16)

p_subtitle = doc.add_paragraph()
p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_subtitle.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
run = p_subtitle.add_run("Differential Gram Staining of a Mixed Smear Containing Escherichia coli and Staphylococcus epidermidis")
run.italic = True; run.font.name = 'Times New Roman'; run.font.size = Pt(13)

for _ in range(3):
    doc.add_paragraph()

for line in [
    "George Vela",
    "BIO203A — Microbiology Laboratory",
    "Spring 2026",
    "Instructor: Dr. Radu Popa",
    "National University, Los Angeles Campus",
    "",
    "Staining Sessions: May 7 and May 9, 2026",
    "Report Submitted: May 2026",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(line)
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SCOPE
# ════════════════════════════════════════════════════════════════════════════
section_heading("Scope")
add_paragraph([
    ("This laboratory exercise performed the Gram stain — the most widely used "
     "differential staining procedure in clinical microbiology — on a single mixed "
     "smear containing two bacterial species with opposite cell-wall types: ", {}),
    ("Escherichia coli", {'italic': True}),
    (" (expected gram-negative) and ", {}),
    ("Staphylococcus epidermidis", {'italic': True}),
    (" (expected gram-positive). The objective was to apply the four-step Gram "
     "stain procedure correctly to a mixed preparation, and to observe both Gram "
     "reactions simultaneously in the same microscope field at oil-immersion "
     "magnification.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
section_heading("Introduction")

add_paragraph([
    ("The Gram stain is a differential staining procedure developed by the Danish "
     "microbiologist Hans Christian Gram in 1884 to distinguish between bacteria "
     "with different types of cell walls. It remains one of the most frequently "
     "used staining techniques in microbiology (Parker et al. 2016, §2.4). Unlike "
     "simple stains, which use a single basic dye to reveal cell shape and "
     "arrangement, differential stains use multiple reagents to give different "
     "groups of bacteria different colors so that they can be distinguished within "
     "the same preparation (Parker et al. 2016, §2.4).", {}),
])

add_paragraph([
    ("The Gram stain procedure consists of four reagents applied in sequence "
     "(Parker et al. 2016, §2.4). First, crystal violet, a primary stain, is "
     "applied to a heat-fixed smear, giving all of the cells a purple color. "
     "Second, Gram's iodine, a mordant, is added; the iodine acts as a trapping "
     "agent that complexes with crystal violet, forming a crystal-violet–iodine "
     "complex that clumps and stays contained in thick layers of peptidoglycan in "
     "the cell wall. Third, a decolorizing agent (ethanol or an acetone/ethanol "
     "solution) is added: cells with thick peptidoglycan retain the crystal "
     "violet–iodine complex and remain purple, whereas cells with thinner "
     "peptidoglycan are decolorized and become colorless. Fourth, a secondary "
     "counterstain — usually safranin — stains the decolorized cells pink while "
     "remaining less noticeable on cells that still contain the crystal violet "
     "dye (Parker et al. 2016, §2.4).", {}),
])

add_paragraph([
    ("The purple, crystal-violet–stained cells are referred to as gram-positive "
     "cells, while the red or pink, safranin-stained cells are gram-negative "
     "(Parker et al. 2016, §2.4). The difference reflects the underlying cell-wall "
     "architecture: gram-positive cell walls have a thick peptidoglycan layer "
     "external to the plasma membrane that retains the crystal-violet–iodine "
     "complex during decolorization, while gram-negative cells have a much thinner "
     "peptidoglycan layer covered by an outer membrane and lose the complex "
     "during the alcohol wash. The two organisms used in this exercise are known "
     "examples of each type. ", {}),
    ("Escherichia coli", {'italic': True}),
    (" is a gram-negative bacillus and member of the family Enterobacteriaceae "
     "within the Gammaproteobacteria (Parker et al. 2016, §4.4). ", {}),
    ("Staphylococcus epidermidis", {'italic': True}),
    (" is a gram-positive coccus belonging to the genus ", {}),
    ("Staphylococcus", {'italic': True}),
    (", whose name derives from the Greek word for bunches of grapes and "
     "describes the characteristic microscopic appearance of its cells in "
     "clusters; ", {}),
    ("Staphylococcus", {'italic': True}),
    (" species are facultative anaerobes, halophilic, and nonmotile, and ", {}),
    ("S. epidermidis", {'italic': True}),
    (" is a common member of the normal microbiota of human skin (Parker et al. "
     "2016, §4.4). Combining one organism of each Gram type on a single mixed "
     "smear allows both reactions — purple and pink — to be observed in one "
     "microscope field, providing direct visual confirmation that the Gram stain "
     "differentiates between cell-wall types.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# MATERIALS AND METHODS
# ════════════════════════════════════════════════════════════════════════════
section_heading("Materials and Methods")

subhead("Materials")
for item in [
    "Bacterial cultures: Escherichia coli and Staphylococcus epidermidis",
    "Clean glass microscope slides",
    "Distilled water in a dropper",
    "Sterile inoculating loops",
    "Bunsen burner with striker",
    "Clothespin (to hold slide during heat fixation)",
    "Crystal violet (primary stain)",
    "Gram's iodine (mordant)",
    "95% ethanol or acetone/ethanol (decolorizing agent)",
    "Safranin (counterstain)",
    "Staining tray, wash bottle of water, and waste beaker",
    "Bibulous paper or paper towel for blotting",
    "Compound brightfield microscope with 100× oil immersion objective",
    "Type-A immersion oil",
    "Lens paper",
]:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(f"• {item}")
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)

subhead("Methods")

add_paragraph([
    ("On May 7, 2026, a mixed smear was prepared on a clean glass slide. A "
     "loopful of distilled water was placed in the center of the slide, and small "
     "amounts of ", {}),
    ("E. coli", {'italic': True}),
    (" and ", {}),
    ("S. epidermidis", {'italic': True}),
    (" were transferred into the water and mixed thoroughly to form a single "
     "smear containing cells of both organisms. The smear was spread evenly to a "
     "thin film, allowed to air-dry completely, and then heat-fixed by passing "
     "the slide through the outer cone of the Bunsen flame two to three times "
     "with the smear side up.", {}),
])

add_paragraph([
    ("The heat-fixed mixed smear was placed on the staining tray and the "
     "four-step Gram stain procedure was applied. Crystal violet was added to "
     "cover the smear and left for one minute, then rinsed gently with distilled "
     "water. Gram's iodine was added next and left for one minute, then rinsed "
     "with distilled water. The slide was tilted and 95% ethanol was applied "
     "drop-by-drop until the solvent ran clear from the smear (approximately "
     "10–20 seconds), and the slide was immediately rinsed with water to halt "
     "decolorization. Safranin was added as a counterstain and left for one "
     "minute, then rinsed and blotted dry with bibulous paper.", {}),
])

add_paragraph([
    ("Staining and observation work continued during the second laboratory "
     "session on May 9, 2026. The stained slide was observed using the compound "
     "brightfield microscope; the smear was first located at 4× and then brought "
     "into focus at 10× and 40× using the parfocal objective progression. A drop "
     "of immersion oil was applied directly to the smear and the 100× oil "
     "immersion objective was rotated into the oil for final observation. A "
     "representative field of the smear was photographed through the ocular.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# RESULTS AND DISCUSSION
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
section_heading("Results and Discussion")

add_paragraph([
    ("The stained mixed smear produced clear differential staining at 1000× "
     "total magnification (Figure 1). Both Gram reactions expected from the "
     "published cell-wall classifications of the two source organisms were "
     "visible in the same microscope field: numerous pink rod-shaped cells "
     "scattered across the field, consistent with the gram-negative bacillus ", {}),
    ("E. coli", {'italic': True}),
    (", together with a smaller number of darker-stained cells consistent with "
     "the gram-positive coccus ", {}),
    ("S. epidermidis", {'italic': True}),
    (". Observation details are summarized in Table 1.", {}),
])

subhead("Mixed Smear at 1000× Total Magnification")
add_image(MIXED_IMG, width_inches=5.0)
add_caption([
    ("Figure 1. ", {'bold': True, 'italic': True}),
    ("Gram stain of a mixed smear containing ", {'italic': True}),
    ("Escherichia coli", {'italic': True}),
    (" and ", {'italic': True}),
    ("Staphylococcus epidermidis", {'italic': True}),
    (" at 1000× total magnification (100× oil immersion × 10× ocular). Numerous "
     "small pink rod-shaped cells are visible across the field, indicating "
     "loss of the crystal-violet–iodine complex during alcohol decolorization "
     "and uptake of the safranin counterstain — the gram-negative reaction "
     "expected of ", {'italic': True}),
    ("E. coli", {'italic': True}),
    (". Darker-stained cells observable in the same field are consistent with "
     "the gram-positive reaction expected of ", {'italic': True}),
    ("S. epidermidis", {'italic': True}),
    (". The black needle visible in the lower portion of the field is the "
     "microscope stage pointer, not part of the specimen.", {'italic': True}),
])

subhead("Summary of Gram Reactions")

table = doc.add_table(rows=3, cols=5)
table.style = 'Table Grid'
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

headers = ['Organism', 'Final Color', 'Gram Reaction', 'Cell Shape', 'Arrangement']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]; cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(h); run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(11)

rows_data = [
    ['Escherichia coli', 'Pink', 'Gram-negative (−)', 'Bacillus (rod)', 'Singles, scattered'],
    ['Staphylococcus epidermidis', 'Purple / darker', 'Gram-positive (+)', 'Coccus (sphere)', 'Clusters'],
]
for i, row in enumerate(rows_data, start=1):
    for j, val in enumerate(row):
        cell = table.rows[i].cells[j]; cell.text = ''
        p = cell.paragraphs[0]; p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(val); run.font.name = 'Times New Roman'; run.font.size = Pt(11)
        if j == 0:
            run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(12)
run = p.add_run("Table 1. ")
run.bold = True; run.italic = True; run.font.name = 'Times New Roman'; run.font.size = Pt(11)
run = p.add_run("Summary of Gram stain observations.")
run.italic = True; run.font.name = 'Times New Roman'; run.font.size = Pt(11)

subhead("Interpretation")
add_paragraph([
    ("The pink final color of the dominant rod-shaped cells in the field is "
     "consistent with the thin peptidoglycan cell wall and outer membrane "
     "characteristic of gram-negative bacteria. During alcohol decolorization, "
     "the crystal-violet–iodine complex is more easily washed out of cells with "
     "thinner peptidoglycan, so these cells become colorless and then take up "
     "the safranin counterstain to appear pink at the end of the procedure "
     "(Parker et al. 2016, §2.4). The rod shape and pink color together are "
     "consistent with the morphology and Gram reaction of ", {}),
    ("E. coli", {'italic': True}),
    (", a member of the family Enterobacteriaceae within the "
     "Gammaproteobacteria — a class of gram-negative bacteria (Parker et al. "
     "2016, §4.4).", {}),
])

add_paragraph([
    ("The darker-stained cells observable in the same field correspond to the "
     "gram-positive component of the mixed smear, consistent with the expected "
     "behavior of ", {}),
    ("S. epidermidis", {'italic': True}),
    (". Gram-positive cells have a thick peptidoglycan layer external to the "
     "plasma membrane that retains the crystal-violet–iodine complex during "
     "alcohol decolorization, so they remain purple at the end of the procedure "
     "(Parker et al. 2016, §2.4). ", {}),
    ("Staphylococcus", {'italic': True}),
    (" species are gram-positive cocci that grow in characteristic clusters "
     "(the genus name derives from the Greek word for bunches of grapes), and ", {}),
    ("S. epidermidis", {'italic': True}),
    (" is a common member of the normal microbiota of human skin (Parker et al. "
     "2016, §4.4).", {}),
])

add_paragraph([
    ("The simultaneous appearance of both pink (gram-negative) and darker, "
     "retained-crystal-violet (gram-positive) cells in the same microscope field "
     "directly demonstrates that the Gram stain reagents acted differentially on "
     "the two cell-wall types, even when the two organisms were mixed on the "
     "same slide. The contrast achieved within one preparation indicates that "
     "the four-step procedure was applied within appropriate time windows: "
     "over-decolorization (excess alcohol exposure) can cause gram-positive "
     "cells to lose the crystal-violet complex and incorrectly appear pink, and "
     "under-decolorization can cause gram-negative cells to retain purple — "
     "neither error appears to have occurred in this preparation.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ════════════════════════════════════════════════════════════════════════════
section_heading("Conclusion")

add_paragraph([
    ("Differential Gram staining of a mixed smear containing ", {}),
    ("Escherichia coli", {'italic': True}),
    (" and ", {}),
    ("Staphylococcus epidermidis", {'italic': True}),
    (" produced both expected Gram reactions visible in the same microscope "
     "field: small pink rod-shaped cells (gram-negative) and darker-stained "
     "cells with retained crystal-violet color (gram-positive). The results are "
     "consistent with the cell-wall classifications of each organism in "
     "OpenStax §4.4 and confirm that the four-step Gram stain procedure "
     "described in §2.4 was carried out correctly.", {}),
])

add_paragraph([
    ("The Gram stain is a foundational tool for the rapid differential "
     "classification of bacterial isolates and underlies subsequent diagnostic "
     "and identification work in clinical microbiology. The same procedure will "
     "be applied to the library-plate isolates from the soil sample in later "
     "exercises of the Antibiotic Discovery Project to begin characterizing the "
     "cell-wall type of each candidate antimicrobial producer.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════════════
section_heading("References")


def add_reference(text_runs):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(0)
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'; run.font.size = Pt(12)


add_reference([
    ("Parker N, Schneegurt M, Tu A-HT, Lister P, Forster BM. 2016. Microbiology. "
     "Houston (TX): OpenStax. Available from: "
     "https://openstax.org/details/books/microbiology", {}),
])

doc.save(OUTPUT)
print(f"Done -> {OUTPUT}")
print(f"Mixed slide photo: {'EMBEDDED' if MIXED_IMG else 'MISSING (placeholder)'}")
if MIXED_IMG: print(f"  {MIXED_IMG}")
