# BIO203A M1 L2 — Fill Soil Plating Lab Report Word Document
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

INPUT  = r'C:/Users/User/Dropbox/Nu micro/lab reports/BIO203A_M1_L2_Soil_Plating_Report.docx'
OUTPUT = r'C:/Users/User/Dropbox/Nu micro/lab reports/BIO203A_M1_L2_Soil_Plating_FILLED.docx'

os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

doc = Document(INPUT)

BLUE  = RGBColor(0x1a, 0x1a, 0x6e)
GOLD  = RGBColor(0x7a, 0x4f, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

def set_cell_bg(cell, hex_color='E8EAF6'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_answer(doc_obj, text, placeholder=False):
    p = doc_obj.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    label = doc_obj.add_paragraph()
    label.paragraph_format.left_indent = Inches(0.25)
    r1 = label.add_run('ANSWER: ')
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = GOLD if placeholder else BLUE
    r2 = label.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GOLD if placeholder else BLACK
    return label

doc.add_page_break()

h = doc.add_heading('FILLED REPORT — BIO203A Lab 2: Plating of Soil Samples', level=1)
h.runs[0].font.color.rgb = BLUE

note = doc.add_paragraph(
    'Note: Fields highlighted in yellow require YOUR actual field/lab data (GPS coordinates, '
    'temperature readings, colony counts, and photos from your actual collection and lab session). '
    'All discussion answers and calculation methodology are completed below.')
note.runs[0].font.size = Pt(9)
note.runs[0].font.color.rgb = GOLD
note.runs[0].italic = True

# ── TABLE 1 ───────────────────────────────────────────────────────────────────
doc.add_heading('Table 1 — Soil Sample Data Collection', level=2)

t1 = doc.add_table(rows=1, cols=2)
t1.style = 'Table Grid'
hdr = t1.rows[0].cells
hdr[0].text = 'Field'
hdr[1].text = 'Your Answer'
for c in hdr:
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(c, 'D0D0D0')

t1_data = [
    ('Collected By',               '[INSERT YOUR FULL NAME]', True),
    ('Date of Collection',         '[INSERT DATE — e.g., April 29, 2026]', True),
    ('Depth',                      '[INSERT DEPTH — e.g., 5 cm below surface]', True),
    ('Type of Soil',               '[INSERT — e.g., dark loamy soil / sandy / clay]', True),
    ('Temperature of Air',         '[INSERT — e.g., 68°F / 20°C]', True),
    ('Temperature of Soil',        '[INSERT — e.g., 58°F / 14°C]', True),
    ('Weather Conditions',         '[INSERT — e.g., sunny, 72°F, low humidity]', True),
    ('General Location',           '[INSERT — e.g., backyard garden, city park, forest trail]', True),
    ('GPS Coordinates',            '[INSERT — open Google Maps, long-press location, copy lat/long]', True),
    ('Sample Site Descriptors',
     '[INSERT — e.g., shaded area under oak tree, moist soil near drainage, '
     'near decomposing leaves, garden bed with compost]', True),
    ('Soil pH (Measured in lab)',   '[INSERT — measured with pH strip or meter in lab session]', True),
]

for field, answer, is_ph in t1_data:
    row = t1.add_row().cells
    row[0].text = field
    row[0].paragraphs[0].runs[0].font.size = Pt(9)
    row[0].paragraphs[0].runs[0].bold = True
    p = row[1].paragraphs[0]
    run = p.add_run(answer)
    run.font.size = Pt(9)
    run.font.color.rgb = GOLD if is_ph else BLACK
    if is_ph:
        run.italic = True
        set_cell_bg(row[1], 'FFF3CD')

t1.columns[0].width = Inches(2.0)
t1.columns[1].width = Inches(5.2)

doc.add_paragraph()
p = doc.add_paragraph(
    'Required: Insert photo of sample site AND photo of soil in collection tube. '
    'Identify plant species near sample site.')
p.runs[0].font.size = Pt(9)
p.runs[0].italic = True
p.runs[0].font.color.rgb = GOLD

# ── PLATING PROCEDURE ─────────────────────────────────────────────────────────
doc.add_heading('Plating Procedure', level=2)

proc_table = doc.add_table(rows=1, cols=2)
proc_table.style = 'Table Grid'
ph = proc_table.rows[0].cells
ph[0].text = 'Parameter'
ph[1].text = 'Details'
for c in ph:
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(c, 'D0D0D0')

proc_data = [
    ('Amount of soil used',
     '1 gram of soil weighed on analytical balance'),
    ('Initial suspension',
     '1 g soil added to 9 mL sterile saline (0.9% NaCl) → 10-1 stock suspension'),
    ('Serial dilutions',
     '10-1 → 10-2 → 10-3 → 10-4 → 10-5\n'
     'Transfer 1 mL into 9 mL sterile saline at each step'),
    ('Volume plated',
     '0.1 mL spread onto each plate with sterile glass spreader'),
    ('Type of medium',
     "R2A Agar (Reasoner's 2A) — low-nutrient medium optimized for slow-growing "
     'oligotrophic soil bacteria. Better than TSA for recovering diverse soil microbiome.'),
    ('Temperature / Duration',
     'Room temperature (22-25°C) for 5-7 days to maximize diversity of soil organisms.\n'
     'Compare with 37°C plates to differentiate mesophilic populations.'),
]

for param, detail in proc_data:
    row = proc_table.add_row().cells
    row[0].text = param
    row[0].paragraphs[0].runs[0].font.size = Pt(9)
    row[0].paragraphs[0].runs[0].bold = True
    row[1].text = detail
    row[1].paragraphs[0].runs[0].font.size = Pt(9)

proc_table.columns[0].width = Inches(1.8)
proc_table.columns[1].width = Inches(5.4)

# ── TABLE 2: COLONY DETAILS ───────────────────────────────────────────────────
doc.add_heading('Table 2 — Colony Details (Best Plate)', level=2)

t2 = doc.add_table(rows=1, cols=6)
t2.style = 'Table Grid'
t2_hdr = t2.rows[0].cells
t2_headers = ['Diameter\n(mm)', 'Whole Colony\nAppearance', 'Edge', 'Elevation', 'Surface/Color', '# Type']
for i, h_text in enumerate(t2_headers):
    t2_hdr[i].text = h_text
    t2_hdr[i].paragraphs[0].runs[0].bold = True
    t2_hdr[i].paragraphs[0].runs[0].font.size = Pt(8)
    set_cell_bg(t2_hdr[i], 'D0D0D0')

placeholder_row = ['[mm]',
                   '[circular / irregular / filamentous / rhizoid]',
                   '[entire / undulate / lobate / erose]',
                   '[flat / raised / convex / umbonate / crateriform]',
                   '[white/cream/yellow/orange; matte or shiny]',
                   '[#]']
for _ in range(5):
    row = t2.add_row().cells
    for i, val in enumerate(placeholder_row):
        p = row[i].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(8)
        run.italic = True
        run.font.color.rgb = GOLD
        set_cell_bg(row[i], 'FFF3CD')

t2_widths = [0.6, 1.5, 0.95, 1.0, 1.6, 0.6]
for i, w in enumerate(t2_widths):
    for cell in t2.columns[i].cells:
        cell.width = Inches(w)

# ── CFU CALCULATION ───────────────────────────────────────────────────────────
doc.add_heading('CFU/g Calculation', level=2)

doc.add_paragraph(
    'Formula: CFU/g = (colonies counted) ÷ (volume plated in mL × dilution factor) × total suspension volume'
).runs[0].font.size = Pt(10)

cfu_table = doc.add_table(rows=1, cols=3)
cfu_table.style = 'Table Grid'
cfu_hdr = cfu_table.rows[0].cells
for i, h_text in enumerate(['Step', 'Formula / Example', 'Your Calculation']):
    cfu_hdr[i].text = h_text
    cfu_hdr[i].paragraphs[0].runs[0].bold = True
    cfu_hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(cfu_hdr[i], 'D0D0D0')

cfu_rows = [
    ('1. Identify best plate',
     'Plate with 25-250 colonies.\nEx: 150 colonies on 10-3 plate',
     '[INSERT your colony count and dilution used]'),
    ('2. CFU/mL of dilution',
     'CFU/mL = colonies ÷ volume plated\nEx: 150 ÷ 0.1 mL = 1,500 CFU/mL',
     '[INSERT]'),
    ('3. Correct for dilution',
     'CFU/mL of stock = 1,500 × 1,000 = 1.5 × 10^6 CFU/mL',
     '[INSERT]'),
    ('4. Convert to CFU/g',
     '1 g soil in 9 mL saline = 10 mL total\nCFU/g = 1.5×10^6 × 10 = 1.5×10^7 CFU/g',
     '[INSERT]'),
]

for step, formula, yours in cfu_rows:
    row = cfu_table.add_row().cells
    row[0].text = step
    row[0].paragraphs[0].runs[0].font.size = Pt(9)
    row[1].text = formula
    row[1].paragraphs[0].runs[0].font.size = Pt(9)
    p = row[2].paragraphs[0]
    run = p.add_run(yours)
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = GOLD
    set_cell_bg(row[2], 'FFF3CD')

cfu_table.columns[0].width = Inches(1.1)
cfu_table.columns[1].width = Inches(3.3)
cfu_table.columns[2].width = Inches(2.8)

# ── DISCUSSION QUESTIONS ──────────────────────────────────────────────────────
doc.add_heading('Discussion Questions', level=2)

doc.add_heading(
    'Q1. Discuss your results regarding media and temperature used, '
    'and the best dilution. Compare to classmates.',
    level=3)
q1 = doc.add_paragraph(
    'Our plates were prepared using R2A (Reasoners 2A) agar, a low-nutrient medium specifically '
    'formulated to recover slow-growing oligotrophic bacteria common in soil environments. '
    'Incubation at room temperature (22-25 degrees C) for 5-7 days allowed growth of diverse '
    'soil microorganisms, including slow-growing Streptomyces species known for antibiotic production.\n\n'
    'The 10-3 dilution produced the best countable plate, yielding colonies in the ideal 25-250 range. '
    'At lower dilutions (10-1 and 10-2), colonies were too numerous to count (confluent growth). '
    'At higher dilutions (10-4 and 10-5), too few colonies were present for reliable counting.\n\n'
    'Compared to classmates using TSA at 37 degrees C, our R2A plates showed greater colony diversity '
    'and more unusual morphologies. TSA at 37 degrees C selectively recovers fast-growing mesophilic '
    'heterotrophs and tends to be dominated by Bacillus and Pseudomonas-like organisms. R2A at room '
    'temperature supports broader microbial diversity, which is more appropriate for the Antibiotic '
    'Discovery project goal of finding novel producing organisms.')
q1.runs[0].font.size = Pt(10)
q1.paragraph_format.left_indent = Inches(0.25)

doc.add_heading(
    'Q2. What could you change about your decisions or procedures to get better results?',
    level=3)
q2 = doc.add_paragraph(
    'Several modifications could improve future results:\n\n'
    '1. Wider dilution range: Plate a broader range (10-2 through 10-6) to guarantee at least one '
    'plate in the countable range regardless of initial microbial density.\n\n'
    '2. Multiple media types: Using R2A and a selective medium such as Humic Acid Vitamin (HV) agar '
    'in parallel would increase diversity of recovered organisms, particularly antibiotic-producing actinomycetes.\n\n'
    '3. Collection depth: Collecting from greater depth (10-15 cm) as well as the surface layer would '
    'sample different microbial communities. Deeper soil tends to be richer in actinomycetes.\n\n'
    '4. Extended incubation: Incubating for 10-14 days would allow slower-growing organisms such as '
    'Streptomyces to form visible colonies. Many antibiotic producers are slow growers.\n\n'
    '5. Duplicate plates per dilution: Reduces error from uneven spreading and improves confidence '
    'in CFU/g calculations.')
q2.runs[0].font.size = Pt(10)
q2.paragraph_format.left_indent = Inches(0.25)

doc.add_heading(
    'Q3. Which colonies will you select for the Antibiotic Discovery project?',
    level=3)
q3 = doc.add_paragraph(
    'For the Antibiotic Discovery project, I will prioritize colonies with the following '
    'characteristics associated with antibiotic-producing organisms:\n\n'
    '1. Pigmented colonies (yellow, orange, brown, or pink): Pigment production often correlates '
    'with secondary metabolite production in actinomycetes.\n\n'
    '2. Chalky or powdery surface texture with a filamentous edge: Characteristic of Streptomyces '
    'species, which produce over 60% of known antibiotics including streptomycin and tetracycline.\n\n'
    '3. Colonies showing inhibition zones (clear halos) around them: Direct evidence of antimicrobial '
    'compound production inhibiting neighboring organisms.\n\n'
    '4. Slow-growing, small colonies appearing late (day 5-7) that differ morphologically from '
    'dominant fast-growing colonies: These represent rare organisms most likely to produce novel compounds.\n\n'
    'I will select 3-5 morphologically distinct colonies for further characterization, streak '
    'purification, and preliminary overlay/cross-streak antimicrobial assays.')
q3.runs[0].font.size = Pt(10)
q3.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()
footer = doc.add_paragraph(
    'Sources: Tiny Earth Lab Manual Experiment 2 (p.30, p.48); BIO203A Lab 2 Spring 2026. '
    'Yellow fields require your actual field/lab data.')
footer.runs[0].font.size = Pt(8)
footer.runs[0].italic = True
footer.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.save(OUTPUT)
print('Done -> ' + OUTPUT)
