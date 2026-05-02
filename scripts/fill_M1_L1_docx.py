# BIO203A M1 L1 — Fill Microscopy Lab Report Word Document
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, copy

INPUT  = r'C:/Users/User/Dropbox/Nu micro/lab reports/BIO203A_M1_L1_Microscopy.docx'
OUTPUT = r'C:/Users/User/Dropbox/Nu micro/lab reports/BIO203A_M1_L1_Microscopy_FILLED.docx'

os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

doc = Document(INPUT)

def add_answer_para(doc_or_parent, text, color=RGBColor(0x1a, 0x1a, 0x6e)):
    """Add a styled answer paragraph."""
    p = doc_or_parent.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run('ANSWER: ')
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p

def set_cell_bg(cell, hex_color='E8EAF6'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_colored_para(doc_obj, text, bg=False):
    p = doc_obj.add_paragraph(text)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x6e)
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    return p

# ── Append filled report to document ─────────────────────────────────────────
doc.add_page_break()

h = doc.add_heading('FILLED REPORT — BIO203A Lab 1: Microscopy', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x6e)

doc.add_heading('Part 1 — Organisms Observed', level=2)

# Organism observation table
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
headers = ['Organism', 'Taxonomic Group / Domain', 'Best Objective\n(Total Mag)', 'Observations', 'Sketch/Photo']
for i, h_text in enumerate(headers):
    hdr_cells[i].text = h_text
    run = hdr_cells[i].paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(9)
    set_cell_bg(hdr_cells[i], 'D0D0D0')

organisms = [
    ('Trypanosoma sp.',
     'Protozoa (Protist)\nUnicellular\nDomain: Eukarya',
     '40× obj. (400× total)\nAlso viewed at\n100× obj. (1000× total)',
     'Prepared blood smear (human tissue sample). '
     'Red blood cells visible as biconcave disc-shaped "donuts" — no nucleus (anucleate). '
     'Trypanosoma trypomastigotes visible between and around RBCs; elongated flagellated '
     'parasites each showing a distinct nucleus and kinetoplast (second dark dot near '
     'flagellum base). Free flagellum clearly visible extending from cell body. '
     'At 40× (400× total): donut-shaped RBCs and elongated parasites distinguishable. '
     'At 100× oil immersion (1000× total): flagellum, nucleus, and kinetoplast of each '
     'parasite resolved in detail. Some host cells appeared lysed with parasites '
     'released into surrounding medium. No photo available.',
     '[No photo taken — sketch: elongated flagellated cell alongside biconcave RBCs]'),
]

for org, tax, mag, obs, sketch in organisms:
    row = table.add_row().cells
    data = [org, tax, mag, obs, sketch]
    for i, d in enumerate(data):
        row[i].text = d
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

# Set column widths
widths = [1.1, 1.4, 1.1, 2.5, 0.6]
for i, width in enumerate(widths):
    for cell in table.columns[i].cells:
        cell.width = Inches(width)

doc.add_paragraph()

# ── Q1 ────────────────────────────────────────────────────────────────────────
doc.add_heading('Q1. Which objective is closest to the slide when in focus?', level=3)
add_answer_para(doc,
    'The 100× oil immersion objective is closest to the slide when in focus. '
    'It has the shortest working distance (~0.1–0.2 mm). '
    'This is why the COARSE adjustment knob must NEVER be used at 40× or 100× — '
    'the objective can crash into and destroy the slide and the lens.')

# ── Q2 ────────────────────────────────────────────────────────────────────────
doc.add_heading('Q2. Correct procedure to put the microscope away:', level=3)
add_answer_para(doc,
    '1. Remove the slide from the stage.\n'
    '2. Clean ALL immersion oil from the 100× objective with lens paper ONLY '
    '(never Kleenex or paper towels — they scratch the lens).\n'
    '3. Clean oil from the stage if present.\n'
    '4. Rotate the nosepiece to the scanning (4×) objective — lowest power, pointed down.\n'
    '5. Lower the stage as far from the objectives as possible using the coarse adjustment knob.\n'
    '6. Leave the power cord to the side — do NOT wrap it around the stage.\n'
    '7. Carry with TWO hands — one on the arm, one under the base — to the storage cabinet.\n'
    '8. Store with the arm pointing outward.')

# ── Q3 ────────────────────────────────────────────────────────────────────────
doc.add_heading(
    'Q3. FOV and cell count calculations — Bacillus (2 µm) and yeast (10 µm), '
    'high-dry (40×) lens, FN = 22 mm', level=3)

calc_table = doc.add_table(rows=1, cols=3)
calc_table.style = 'Table Grid'
h_cells = calc_table.rows[0].cells
for i, h_text in enumerate(['Step', 'Calculation', 'Result']):
    h_cells[i].text = h_text
    h_cells[i].paragraphs[0].runs[0].bold = True
    h_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(h_cells[i], 'D0D0D0')

calc_rows = [
    ('Total magnification', '40× (objective) × 10× (ocular)', '400×'),
    ('FOV diameter', 'FN ÷ Total magnification = 22 mm ÷ 400', '0.055 mm = 55 µm'),
    ('Bacillus cells across FOV\n(cell length = 2 µm)', '55 µm ÷ 2 µm per cell', '≈ 27 cells'),
    ('Yeast cells across FOV\n(cell length = 10 µm)', '55 µm ÷ 10 µm per cell', '≈ 5–6 cells'),
]
for step, calc, result in calc_rows:
    row = calc_table.add_row().cells
    for i, d in enumerate([step, calc, result]):
        row[i].text = d
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# ── Inline procedure questions ────────────────────────────────────────────────
doc.add_heading('Procedure Questions (from Lab Steps)', level=2)

doc.add_heading(
    'Step 7: To use a minimum of light, which part should be adjusted?', level=3)
add_answer_para(doc,
    'Close the iris diaphragm (part of the condenser assembly). '
    'This reduces the cone of light entering the condenser and objective lens, '
    'which increases contrast when viewing unstained or low-contrast specimens at low magnification.')

doc.add_heading(
    'Step 8a: When switching from low power to high-dry, only a slight adjustment is needed. Why?', level=3)
add_answer_para(doc,
    'The objective lenses are parfocal — designed so that when the specimen is in focus '
    'at one objective, switching to another objective places it at approximately the same '
    'focal distance. Only minor fine-focus adjustment is needed after switching. '
    '(Popa Ch.2 slide 9 — parfocality is specifically tested.)')

doc.add_heading(
    'Step 8b: More light is usually needed at higher magnification. How can you increase it?', level=3)
add_answer_para(doc,
    '1. Open the iris diaphragm wider — allows more light into the condenser.\n'
    '2. Raise the condenser closer to the slide — concentrates more light onto the specimen.\n'
    '3. Increase lamp voltage/brightness using the illuminator control on the base.\n'
    'At 100× (oil immersion): immersion oil also increases light entering the objective '
    'by preventing refraction at the glass-air interface.')

doc.add_heading(
    'Step 10: Did color change with different lenses? Did FOV change? '
    'How many cells fit at higher vs lower magnification?', level=3)
add_answer_para(doc,
    'Color: Color does not change when switching objectives — the same light source and '
    'stain are used. However, brightness decreases at higher magnification (more light needed).\n\n'
    'FOV: Yes — FOV decreases as magnification increases. Fewer cells visible at 1000× than at 100×.\n\n'
    'Cell count: At 100× total (FOV ≈ 220 µm) many cells span the field. '
    'At 1000× total (FOV ≈ 22 µm) only 1–3 Bacillus cells may span the field.')

doc.add_paragraph()
p = doc.add_paragraph(
    'Sources: BIO203A Lab 1 Microscopy (Popa); Popa Ch.2 lecture slides. Spring 2026. '
    'Diagrams and cell phone photos must be added from your own in-lab observations.')
p.runs[0].font.size = Pt(8)
p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
p.runs[0].italic = True

doc.save(OUTPUT)
print('Done -> ' + OUTPUT)
