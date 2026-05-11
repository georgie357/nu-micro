# BIO203A Lab 4: Aseptic Technique — Traditional Lab Report (Popa format, STRICT v2)
# Author: George Vela
# References cite ONLY OpenStax Microbiology (Parker et al. 2016). CSE Name-Year style.

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\User\Dropbox\Nu micro\lab reports\BIO203A_Lab4_Traditional_Report.docx"

IMG_BACILLUS_SLANT = r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 4 bacillus slant.JPG"
IMG_BACILLUS_DEEP = r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 4 bacillus deep.JPG"
IMG_ECOLI_SLANT = r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 4 ecoli slant.JPG"
IMG_ECOLI_DEEP = r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 4 ecoli deep.JPG"

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
style.paragraph_format.space_after = Pt(0)

for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def add_page_number(footer_para):
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


for section in doc.sections:
    add_page_number(section.footer.paragraphs[0])


def add_paragraph(text_runs, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_indent=Inches(0.5), space_after=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.first_line_indent = first_indent
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = space_after
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(fmt.get('size', 12))
        if fmt.get('bold'): run.bold = True
        if fmt.get('italic'): run.italic = True


def section_heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text); run.bold = True
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)


def subhead(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text); run.bold = True; run.italic = True
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)


def add_caption(text_runs):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(12)
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'; run.font.size = Pt(11)
        if fmt.get('bold'): run.bold = True
        if fmt.get('italic', True): run.italic = True


def add_image(path, width_inches=3.5):
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))


# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
run = p_title.add_run("Lab 4: Aseptic Technique")
run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(16)

p_subtitle = doc.add_paragraph()
p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_subtitle.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
run = p_subtitle.add_run("Aseptic Inoculation of Bacillus subtilis and Escherichia coli on TSA Slant and Deep Culture Tubes")
run.italic = True; run.font.name = 'Times New Roman'; run.font.size = Pt(13)

for _ in range(3):
    doc.add_paragraph()

for line in [
    "George Vela",
    "BIO203A — Microbiology Laboratory",
    "Spring 2026",
    "Instructor: Dr. Radu Popa",
    "National University, Los Angeles Campus",
    "",
    "Inoculation Date: May 5, 2026",
    "Report Submitted: May 2026",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(line)
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SCOPE
# ════════════════════════════════════════════════════════════════════════════
section_heading("Scope")
add_paragraph([
    ("This laboratory exercise practiced the aseptic transfer of bacterial cultures "
     "from pure stock into fresh culture media of two different physical forms "
     "(agar slants and agar deeps) and recorded the resulting growth patterns after "
     "incubation. Two organisms with different oxygen requirements were used — ", {}),
    ("Bacillus subtilis", {'italic': True}),
    (" and ", {}),
    ("Escherichia coli", {'italic': True}),
    (" — so that the relationship between an organism's oxygen requirements and its "
     "growth distribution in a tube culture could be observed directly.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
section_heading("Introduction")

add_paragraph([
    ("Aseptic technique is the set of laboratory practices used to prevent "
     "contamination of sterile surfaces and culture media by unwanted "
     "microorganisms (Parker et al. 2016, §13.1). Sterilizing the inoculating loop "
     "by passing it through a flame — \"flaming the loop\" — is a standard component "
     "of aseptic technique in the microbiology laboratory and ensures that no live "
     "organisms are carried from one culture into another (Parker et al. 2016, §13.2).", {}),
])

add_paragraph([
    ("The two organisms used in this exercise differ in their oxygen requirements. "
     "Members of the genus ", {}),
    ("Bacillus", {'italic': True}),
    (" are gram-positive, large bacillus-shaped bacteria that include aerobes or "
     "facultative anaerobes and can produce endospores (Parker et al. 2016, §4.4). ", {}),
    ("Escherichia coli", {'italic': True}),
    (" is a member of the family Enterobacteriaceae within the Gammaproteobacteria, "
     "the largest and most diverse class of bacteria; ", {}),
    ("E. coli", {'italic': True}),
    (" is found in the human gut where it forms mutualistic relationships and "
     "produces vitamin K, and is the most-studied bacterium of all (Parker et al. "
     "2016, §4.4). The Enterobacteriaceae as a group are facultative anaerobes "
     "(Parker et al. 2016, §9.2).", {}),
])

add_paragraph([
    ("Oxygen requirements of microorganisms can be observed directly by inoculating "
     "the organism into a tube culture in which oxygen concentration decreases with "
     "depth. In such tubes, obligate aerobes grow only near the surface, where "
     "oxygen is abundant; facultative anaerobes are able to grow throughout the "
     "tube, with denser growth often visible at the surface where energy yield from "
     "aerobic respiration is highest (Parker et al. 2016, §9.2). Comparing the "
     "distribution of growth along the stab line of a deep tube inoculated with a "
     "given organism therefore provides a quick visual classification of its oxygen "
     "requirements.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# MATERIALS AND METHODS
# ════════════════════════════════════════════════════════════════════════════
section_heading("Materials and Methods")

subhead("Materials")
for item in [
    "Pure stock culture of Bacillus subtilis",
    "Pure stock culture of Escherichia coli",
    "Tryptic Soy Agar (TSA) slant tubes, sterile (×2)",
    "Tryptic Soy Agar (TSA) deep tubes, sterile (×2)",
    "Inoculating loop (for slant inoculation)",
    "Inoculating needle (for deep stab inoculation)",
    "Bunsen burner with striker",
    "Test-tube rack",
    "Permanent marker for tube labeling",
    "Incubator (~25 °C, room temperature)",
]:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(f"• {item}")
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)

subhead("Methods")

add_paragraph([
    ("On May 5, 2026, four sterile TSA tubes were prepared for inoculation: two "
     "slant tubes and two deep tubes. Each tube was labeled on the side with the "
     "investigator's initials, the date, the medium (TSA), and the organism that "
     "would be inoculated into it. The four tubes were placed in a test-tube rack "
     "alongside the two pure stock cultures.", {}),
])

add_paragraph([
    ("The Bunsen burner was lit and the inoculating loop was sterilized by holding "
     "it in the inner cone of the flame until it glowed red, then briefly cooled in "
     "the air near the flame for approximately 30 seconds. The cap of the ", {}),
    ("B. subtilis", {'italic': True}),
    (" stock tube was removed with the small finger of the dominant hand, the mouth "
     "of the stock tube was flamed, and a loopful of culture was withdrawn. The "
     "stock tube was re-flamed at the mouth, recapped, and returned to the rack. "
     "The cap of the fresh ", {}),
    ("B. subtilis", {'italic': True}),
    (" slant tube was removed, its mouth was flamed, and the loop was used to streak "
     "the inoculum in a zigzag pattern across the agar surface from the bottom of "
     "the slant to the top. The slant tube was flamed at the mouth, recapped, and "
     "returned to the rack. The inoculating loop was flamed thoroughly before being "
     "set down.", {}),
])

add_paragraph([
    ("The deep tube of ", {}),
    ("B. subtilis", {'italic': True}),
    (" was inoculated using an inoculating needle rather than a loop, following the "
     "same aseptic flaming-and-capping protocol described above. A small amount of "
     "culture was picked up on the tip of the sterile needle and the needle was "
     "inserted vertically into the center of the agar deep, stabbed almost to the "
     "bottom of the tube, and withdrawn along the same path. The procedure was then "
     "repeated for the ", {}),
    ("E. coli", {'italic': True}),
    (" stock culture, producing one TSA slant and one TSA deep inoculated with ", {}),
    ("E. coli", {'italic': True}),
    (". All four inoculated tubes were capped loosely to allow gas exchange, "
     "incubated upright at room temperature (~25 °C), and observed for growth.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# RESULTS AND DISCUSSION
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
section_heading("Results and Discussion")

add_paragraph([
    ("All four inoculated tubes showed visible bacterial growth after incubation. "
     "Slant tubes (Figures 1 and 3) developed colonies along the streak line on the "
     "agar surface, and deep tubes (Figures 2 and 4) developed growth distributed "
     "differently along the stab line for the two organisms. The patterns observed "
     "are summarized in Table 1.", {}),
])

# Images — 2x2 layout
subhead("Slant Cultures")
add_image(IMG_BACILLUS_SLANT, width_inches=4.0)
add_caption([
    ("Figure 1. ", {'bold': True, 'italic': True}),
    ("TSA slant tube inoculated with ", {'italic': True}),
    ("Bacillus subtilis", {'italic': True}),
    (" on 5/5/2026. Heavy white/cream growth is visible along the streak line on "
     "the slant surface, with a slightly raised, spreading texture.", {'italic': True}),
])

add_image(IMG_ECOLI_SLANT, width_inches=4.0)
add_caption([
    ("Figure 2. ", {'bold': True, 'italic': True}),
    ("TSA slant tube inoculated with ", {'italic': True}),
    ("Escherichia coli", {'italic': True}),
    (" on 5/5/2026. Moderate white/cream growth is visible along the streak line "
     "with a relatively even, filiform-to-spreading distribution.", {'italic': True}),
])

subhead("Deep Cultures")
add_image(IMG_BACILLUS_DEEP, width_inches=4.0)
add_caption([
    ("Figure 3. ", {'bold': True, 'italic': True}),
    ("TSA deep tube inoculated with ", {'italic': True}),
    ("Bacillus subtilis", {'italic': True}),
    (" on 5/5/2026. Growth is concentrated at the top of the stab line, near the "
     "agar surface where oxygen is most abundant, consistent with aerobic "
     "metabolism.", {'italic': True}),
])

add_image(IMG_ECOLI_DEEP, width_inches=4.0)
add_caption([
    ("Figure 4. ", {'bold': True, 'italic': True}),
    ("TSA deep tube inoculated with ", {'italic': True}),
    ("Escherichia coli", {'italic': True}),
    (" on 5/5/2026. Growth is distributed throughout the length of the stab line, "
     "from the surface to the bottom of the tube, consistent with facultative "
     "anaerobic metabolism.", {'italic': True}),
])

subhead("Summary of Growth Patterns")

# Table
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

headers = ['Tube', 'Organism', 'Growth Pattern', 'Color / Texture']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]; cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(h); run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(11)

rows_data = [
    ['Slant', 'Bacillus subtilis', 'Heavy growth across streak; spreading along slant surface', 'White / cream, slightly raised'],
    ['Slant', 'Escherichia coli', 'Moderate growth along streak; filiform-to-spreading pattern', 'White / cream, smooth'],
    ['Deep', 'Bacillus subtilis', 'Growth concentrated near top of stab line (aerobic zone)', 'Cream, dense at surface'],
    ['Deep', 'Escherichia coli', 'Growth distributed throughout stab line, top to bottom', 'Cream, even distribution'],
]
italic_cells = {(1, 1), (2, 1), (3, 1), (4, 1)}
for i, row in enumerate(rows_data, start=1):
    for j, val in enumerate(row):
        cell = table.rows[i].cells[j]; cell.text = ''
        p = cell.paragraphs[0]; p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(val); run.font.name = 'Times New Roman'; run.font.size = Pt(11)
        if (i, j) in italic_cells:
            run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(12)
run = p.add_run("Table 1. ")
run.bold = True; run.italic = True; run.font.name = 'Times New Roman'; run.font.size = Pt(11)
run = p.add_run("Summary of growth patterns observed on TSA slant and deep tubes after incubation at room temperature.")
run.italic = True; run.font.name = 'Times New Roman'; run.font.size = Pt(11)

subhead("Interpretation")
add_paragraph([
    ("The slant cultures of both organisms produced visible growth along the streak "
     "line, confirming that aseptic transfer was successful and that the organisms "
     "remained viable through the inoculation procedure. The slight differences in "
     "surface texture between the two slants — ", {}),
    ("B. subtilis", {'italic': True}),
    (" producing a heavier, more raised colony than ", {}),
    ("E. coli", {'italic': True}),
    (" — are consistent with morphological descriptions of these two organisms; ", {}),
    ("Bacillus", {'italic': True}),
    (" is a large gram-positive bacillus, whereas ", {}),
    ("E. coli", {'italic': True}),
    (" is a smaller gram-negative bacillus in the Enterobacteriaceae (Parker et al. "
     "2016, §4.4).", {}),
])

add_paragraph([
    ("The deep tube results are the most informative for classifying oxygen "
     "requirement. In the ", {}),
    ("B. subtilis", {'italic': True}),
    (" deep, growth was concentrated at the top of the stab line, near the surface "
     "of the agar where oxygen diffusion is highest, with little or no growth in "
     "the lower portion of the tube. This distribution is consistent with aerobic "
     "metabolism: obligate aerobes are confined to the oxygen-rich surface of a "
     "tube culture (Parker et al. 2016, §9.2). The genus ", {}),
    ("Bacillus", {'italic': True}),
    (" is classified as aerobic or facultatively anaerobic (Parker et al. 2016, "
     "§4.4); the ", {}),
    ("B. subtilis", {'italic': True}),
    (" deep pattern observed in this experiment is consistent with the aerobic end "
     "of that range.", {}),
])

add_paragraph([
    ("In the ", {}),
    ("E. coli", {'italic': True}),
    (" deep, by contrast, growth was distributed along the entire length of the "
     "stab line from the surface to the bottom of the tube. This pattern is "
     "characteristic of facultative anaerobes, organisms that can grow with or "
     "without oxygen and therefore are not restricted to the surface of the medium "
     "(Parker et al. 2016, §9.2). Members of the Enterobacteriaceae, including ", {}),
    ("E. coli", {'italic': True}),
    (", are facultative anaerobes, and the observed deep-tube growth pattern is "
     "consistent with that classification (Parker et al. 2016, §4.4 and §9.2).", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ════════════════════════════════════════════════════════════════════════════
section_heading("Conclusion")

add_paragraph([
    ("Aseptic transfer of pure cultures of ", {}),
    ("Bacillus subtilis", {'italic': True}),
    (" and ", {}),
    ("Escherichia coli", {'italic': True}),
    (" onto fresh TSA slant and deep tubes was performed successfully. After "
     "incubation, all four tubes showed visible growth without evidence of "
     "contamination, confirming that the flaming, capping, and transfer protocols "
     "of aseptic technique were carried out correctly (Parker et al. 2016, §13.1 "
     "and §13.2).", {}),
])

add_paragraph([
    ("The deep-tube growth distributions observed for the two organisms directly "
     "demonstrated their textbook-described oxygen requirements: ", {}),
    ("B. subtilis", {'italic': True}),
    (" growth was concentrated at the oxygen-rich top of the stab line (aerobic "
     "pattern), while ", {}),
    ("E. coli", {'italic': True}),
    (" growth was distributed along the entire stab line (facultative anaerobe "
     "pattern), in agreement with their classifications in OpenStax §4.4 and §9.2.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════════════
section_heading("References")


def add_reference(text_runs):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(0)
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'; run.font.size = Pt(12)


add_reference([
    ("Parker N, Schneegurt M, Tu A-HT, Lister P, Forster BM. 2016. Microbiology. "
     "Houston (TX): OpenStax. Available from: "
     "https://openstax.org/details/books/microbiology", {}),
])

doc.save(OUTPUT)
print(f"Done -> {OUTPUT}")
for label, p in [('Bacillus slant', IMG_BACILLUS_SLANT), ('Bacillus deep', IMG_BACILLUS_DEEP),
                  ('E. coli slant', IMG_ECOLI_SLANT), ('E. coli deep', IMG_ECOLI_DEEP)]:
    print(f"  {label}: {'EMBEDDED' if os.path.exists(p) else 'MISSING'} ({p})")
