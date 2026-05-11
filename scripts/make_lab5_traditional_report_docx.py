# BIO203A Lab 5: Smears and Simple Staining — Traditional Lab Report (Popa STRICT v2)
# Author: George Vela
# References cite ONLY OpenStax Microbiology (Parker et al. 2016). CSE Name-Year style.

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\User\Dropbox\Nu micro\lab reports\BIO203A_Lab5_Traditional_Report.docx"
STAIN_IMG_PATHS = [
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 5 simple stain.JPG",
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 5 simple stain.jpg",
    r"C:\Users\User\Dropbox\Nu micro\lab reports\lab 5 simple stain.png",
]
STAIN_IMG = next((p for p in STAIN_IMG_PATHS if os.path.exists(p)), None)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
style.paragraph_format.space_after = Pt(0)

for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def add_page_number(footer_para):
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


for section in doc.sections:
    add_page_number(section.footer.paragraphs[0])


def add_paragraph(text_runs, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_indent=Inches(0.5), space_after=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.first_line_indent = first_indent
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = space_after
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'; run.font.size = Pt(fmt.get('size', 12))
        if fmt.get('bold'): run.bold = True
        if fmt.get('italic'): run.italic = True


def section_heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text); run.bold = True
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)


def subhead(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text); run.bold = True; run.italic = True
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)


def add_caption(text_runs):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(12)
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'; run.font.size = Pt(11)
        if fmt.get('bold'): run.bold = True
        if fmt.get('italic', True): run.italic = True


def add_image(path, width_inches=4.5):
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))


# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
run = p_title.add_run("Lab 5: Smears and Simple Staining")
run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(16)

p_subtitle = doc.add_paragraph()
p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_subtitle.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
run = p_subtitle.add_run("Preparation, Heat-Fixation, and Methylene Blue Simple Staining of a Bacterial Smear")
run.italic = True; run.font.name = 'Times New Roman'; run.font.size = Pt(13)

for _ in range(3):
    doc.add_paragraph()

for line in [
    "George Vela",
    "BIO203A — Microbiology Laboratory",
    "Spring 2026",
    "Instructor: Dr. Radu Popa",
    "National University, Los Angeles Campus",
    "",
    "Staining Date: May 7, 2026",
    "Report Submitted: May 2026",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(line)
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SCOPE
# ════════════════════════════════════════════════════════════════════════════
section_heading("Scope")
add_paragraph([
    ("This laboratory exercise prepared, heat-fixed, and stained a bacterial smear "
     "using a single basic dye (methylene blue) and observed the resulting "
     "preparation under the compound brightfield microscope at oil immersion. The "
     "purpose was to demonstrate the simple-stain technique and to describe the "
     "morphology — shape and arrangement — of the cells in the stained smear.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
section_heading("Introduction")

add_paragraph([
    ("Most bacterial cells are nearly colorless and have very little contrast against "
     "the surrounding medium when viewed unstained under a brightfield microscope. "
     "Stains are added to a specimen to increase contrast and to make particular "
     "structures visible (Parker et al. 2016, §2.4). In a stain, the colored ion of "
     "the dye is called the chromophore; if the chromophore is the positively "
     "charged ion, the stain is classified as a basic dye, and if the negative ion "
     "is the chromophore, the stain is classified as an acidic dye (Parker et al. "
     "2016, §2.4).", {}),
])

add_paragraph([
    ("Bacterial cell walls typically carry a net negative charge, so the positively "
     "charged chromophores of basic dyes are attracted to and stick to the cell "
     "walls; this makes basic dyes function as positive stains in which the cells "
     "themselves take up the color (Parker et al. 2016, §2.4). Commonly used basic "
     "dyes include crystal violet, malachite green, methylene blue, and safranin "
     "(Parker et al. 2016, §2.4). In simple staining, a single dye is applied to a "
     "fixed smear to emphasize particular structures of the specimen, most often "
     "cell shape and arrangement (Parker et al. 2016, §2.4).", {}),
])

add_paragraph([
    ("Prokaryotic cells of the same species are typically similar in shape, or "
     "cell morphology, and cells may also group together in characteristic "
     "arrangements (Parker et al. 2016, §3.3). The three most common bacterial cell "
     "shapes are the coccus (sphere), the bacillus (rod), and the spiral, and "
     "arrangements such as pairs, chains, and clusters arise from the planes in "
     "which the cells divide and remain attached to each other (Parker et al. 2016, "
     "§3.3). A simple stain followed by examination at high magnification — most "
     "commonly the 100× oil immersion objective — allows direct visualization of "
     "both characteristics.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# MATERIALS AND METHODS
# ════════════════════════════════════════════════════════════════════════════
section_heading("Materials and Methods")

subhead("Materials")
for item in [
    "Bacterial culture (slant or broth) — source culture used for the smear",
    "Clean glass microscope slide",
    "Distilled water (for solid culture smears) in a dropper",
    "Sterile inoculating loop",
    "Bunsen burner with striker",
    "Clothespin (to hold slide during heat fixation)",
    "Permanent marker or pencil for labeling",
    "Methylene blue stain (basic dye)",
    "Staining tray, wash bottle of water, and beaker for waste",
    "Bibulous (blotting) paper or paper towel",
    "Compound brightfield microscope with 4×, 10×, 40×, and 100× oil immersion objectives",
    "Type-A immersion oil",
    "Lens paper",
]:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(f"• {item}")
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)

subhead("Methods")

add_paragraph([
    ("On May 7, 2026, a clean glass slide was labeled with the investigator's "
     "initials and the date in pencil on the frosted end. A circle was drawn on "
     "the underside of the slide with a permanent marker to indicate the location "
     "of the smear. The inoculating loop was sterilized by flaming until red and "
     "cooled briefly in the air near the flame. A loopful of culture was applied "
     "to the circled area of the slide and spread thinly and evenly to produce a "
     "smear approximately the size of the marker circle. The loop was flamed again "
     "before being set down.", {}),
])

add_paragraph([
    ("The slide was allowed to air-dry completely on the bench top. The dried "
     "smear was then heat-fixed by holding the slide with a clothespin and passing "
     "it through the outer blue cone of the Bunsen flame two to three times, with "
     "the smear-side up. The slide was checked to make sure it was warm but not "
     "hot enough to damage the cells, and then placed on the staining tray.", {}),
])

add_paragraph([
    ("Methylene blue solution was applied to cover the smear circle and was left "
     "in contact with the smear for one minute. After one minute, the slide was "
     "rinsed gently with distilled water from a wash bottle, with the slide tilted "
     "so that the rinse water drained away from the smear and into the waste "
     "beaker. The slide was blotted dry by pressing the edge into bibulous paper "
     "and allowing the smear area itself to air-dry briefly before observation.", {}),
])

add_paragraph([
    ("The dried, stained slide was placed on the microscope stage and the smear "
     "was first located at 4× and then brought into focus at 10× and 40× using "
     "the parfocal objective progression. The 40× objective was rotated out of "
     "the optical path, a single drop of immersion oil was applied directly to "
     "the smear circle on the slide, and the 100× oil immersion objective was "
     "rotated into the oil. Fine focus was used to bring the cells into sharp "
     "resolution and a representative field was photographed through the ocular.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# RESULTS AND DISCUSSION
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
section_heading("Results and Discussion")

subhead("Stained Smear at 1000× Total Magnification")

add_image(STAIN_IMG, width_inches=4.5)
add_caption([
    ("Figure 1. ", {'bold': True, 'italic': True}),
    ("Photograph through the ocular of the methylene blue simple stain at 1000× "
     "total magnification (100× oil immersion objective × 10× ocular). Numerous "
     "small rod-shaped cells stained blue are visible across the field, with "
     "scattered single cells in clearer regions and denser aggregates of cells in "
     "the darker-stained patches. The dark, needle-like feature in the lower-left "
     "portion of the field is the stage pointer of the microscope, not part of "
     "the specimen.", {'italic': True}),
])

subhead("Description of Observed Cells")
add_paragraph([
    ("The smear contained large numbers of small, rod-shaped cells that took up "
     "the methylene blue stain and appeared blue against the lighter background. "
     "Individual cells were visible as short, straight rods of consistent size and "
     "shape, with no obvious chains or rosettes; cells in the densely stained "
     "regions of the field appeared aggregated rather than organized into "
     "distinctive multi-cell arrangements. No cocci, spirals, or branching "
     "filaments were observed.", {}),
])

add_paragraph([
    ("The uptake of methylene blue is consistent with the expected behavior of a "
     "basic dye on bacterial cells. The positively charged methylene blue "
     "chromophore is attracted to the negatively charged components of the "
     "bacterial cell wall, so the cells themselves take up the color while the "
     "surrounding medium remains pale; this is the defining feature of a positive "
     "(direct) simple stain (Parker et al. 2016, §2.4).", {}),
])

subhead("Identification of the Stained Organism")
add_paragraph([
    ("The cells observed in this preparation were small, straight, rod-shaped "
     "bacilli — a morphology consistent with the gram-negative enteric bacterium ", {}),
    ("Escherichia coli", {'italic': True}),
    (", which is a member of the family Enterobacteriaceae within the "
     "Gammaproteobacteria (Parker et al. 2016, §4.4). However, species-level "
     "identification cannot be confirmed from a simple stain alone, because a "
     "simple stain reveals only cell shape and arrangement and not the differential "
     "characteristics (such as cell-wall type or biochemical activity) that are "
     "needed for definitive identification. The morphology observed is consistent "
     "with the source culture being ", {}),
    ("E. coli", {'italic': True}),
    (", but additional staining or testing would be required to confirm.", {}),
])

# Summary table
subhead("Summary of Observation")
table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

rows_data = [
    ['Feature', 'Observation'],
    ['Stain used', 'Methylene blue (basic dye, positive simple stain)'],
    ['Fixation', 'Heat-fixed by passing slide through outer cone of Bunsen flame'],
    ['Objective / magnification', '100× oil immersion / 1000× total magnification'],
    ['Cell shape', 'Bacilli (small straight rods)'],
    ['Arrangement', 'Singles with regions of aggregated cells; no chains, clusters, or pairs distinctly resolved'],
]
for i, row in enumerate(rows_data):
    for j, val in enumerate(row):
        cell = table.rows[i].cells[j]; cell.text = ''
        p = cell.paragraphs[0]; p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(val); run.font.name = 'Times New Roman'; run.font.size = Pt(11)
        if i == 0:
            run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(12)
run = p.add_run("Table 1. ")
run.bold = True; run.italic = True; run.font.name = 'Times New Roman'; run.font.size = Pt(11)
run = p.add_run("Summary of the methylene blue simple stain observation.")
run.italic = True; run.font.name = 'Times New Roman'; run.font.size = Pt(11)

# ════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ════════════════════════════════════════════════════════════════════════════
section_heading("Conclusion")

add_paragraph([
    ("A bacterial smear was prepared, heat-fixed, and stained with methylene blue, "
     "and the resulting preparation was observed at 1000× total magnification under "
     "oil immersion. The methylene blue functioned as a positive simple stain "
     "consistent with its classification as a basic dye whose positively charged "
     "chromophore binds to the negatively charged bacterial cell wall (Parker et "
     "al. 2016, §2.4). The cells in the smear were small, straight bacilli, "
     "consistent with the rod-shape category described in OpenStax §3.3, and the "
     "morphology was compatible with — though not diagnostic for — ", {}),
    ("Escherichia coli", {'italic': True}),
    (" (Parker et al. 2016, §4.4).", {}),
])

add_paragraph([
    ("The exercise reinforced the role of the simple stain as a quick first-pass "
     "tool for describing the shape and arrangement of cells in a smear, and "
     "highlighted its limitation: simple stains do not differentiate among "
     "bacterial groups, and a differential stain (such as the Gram stain) is "
     "needed for taxonomic discrimination beyond morphology.", {}),
])

# ════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════════════
section_heading("References")


def add_reference(text_runs):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(0)
    for text, fmt in text_runs:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'; run.font.size = Pt(12)


add_reference([
    ("Parker N, Schneegurt M, Tu A-HT, Lister P, Forster BM. 2016. Microbiology. "
     "Houston (TX): OpenStax. Available from: "
     "https://openstax.org/details/books/microbiology", {}),
])

doc.save(OUTPUT)
print(f"Done -> {OUTPUT}")
print(f"Stain photo: {'EMBEDDED' if STAIN_IMG else 'MISSING'} ({STAIN_IMG})")
